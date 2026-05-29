"""
Hire Man — Technical Audit Report Generator
Produces a fully-formatted Word (.docx) document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
BRAND_BLUE  = RGBColor(0x1A, 0x73, 0xE8)   # #1A73E8
DARK_GREY   = RGBColor(0x20, 0x20, 0x20)
MID_GREY    = RGBColor(0x55, 0x55, 0x55)
RED_FLAG    = RGBColor(0xC0, 0x39, 0x2B)
AMBER_FLAG  = RGBColor(0xD3, 0x7B, 0x00)
GREEN_OK    = RGBColor(0x1E, 0x8E, 0x3E)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

# ─────────────────────────────────────────────────────────────────────────────
# Helper functions
# ─────────────────────────────────────────────────────────────────────────────

def shade_cell(cell, hex_color: str):
    """Fill a table cell with a background colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_col_widths(table, widths):
    """Set column widths (Inches) for a table."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])

def add_heading(level: int, text: str):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BRAND_BLUE if level == 1 else DARK_GREY
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
    return p

def add_para(text: str, bold=False, color=None, size=10, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK_GREY
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(text: str, color=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + indent_level * 0.2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = DARK_GREY

def add_warning(text: str):
    """Add a red-highlighted warning paragraph."""
    p = doc.add_paragraph()
    run = p.add_run('⚠  ' + text)
    run.font.color.rgb = RED_FLAG
    run.font.bold = True
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_note(text: str):
    """Add an amber/info note paragraph."""
    p = doc.add_paragraph()
    run = p.add_run('ℹ  ' + text)
    run.font.color.rgb = AMBER_FLAG
    run.font.bold = False
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def simple_table(headers, rows, col_widths=None):
    """Create a simple bordered table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr_cells[i], '1A73E8')
        run = hdr_cells[i].paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.color.rgb = DARK_GREY

    if col_widths:
        set_col_widths(table, col_widths)

    doc.add_paragraph()  # spacing after table
    return table

# ─────────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('HIRE')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = BRAND_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PRODUCT TECHNICAL AUDIT REPORT')
r.font.size  = Pt(16)
r.font.bold  = True
r.font.color.rgb = DARK_GREY

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Generated: {datetime.date.today().strftime("%d %B %Y")}')
r.font.size  = Pt(11)
r.font.color.rgb = MID_GREY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Confidential — Internal Use Only')
r.font.size  = Pt(10)
r.font.italic = True
r.font.color.rgb = MID_GREY

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — PRODUCT IDENTITY
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '1. Product Identity')

simple_table(
    ['Attribute', 'Value'],
    [
        ['Product Name',       'Hire'],
        ['Tagline',            '"Your Career. On Autopilot."'],
        ['Primary Purpose',    'AI-powered career development SaaS platform for job seekers and employers — covering resume building, AI job matching, autonomous job applications, salary negotiation coaching, interview intelligence, and employer ATS/hiring pipeline management.'],
        ['Framework',          'Laravel 12.x'],
        ['Language / Runtime', 'PHP 8.2+'],
        ['Frontend Type',      'Blade + Livewire 3.x (server-side reactive components) + Alpine.js + Tailwind CSS 3.x'],
        ['Admin UI',           'Filament 4.x (Filament PHP) — panel at /studai'],
        ['Build Tool',         'Vite 7.x'],
        ['Deployment',         'Azure App Service (Linux container) — https://studai-app-prod.azurewebsites.net'],
        ['Frontend ↔ Backend', '(1) Standard HTML form POSTs for web pages  (2) Livewire AJAX wire calls for reactive components  (3) Sanctum-authenticated JSON REST API for mobile/SPA consumption  (4) Custom API token (ApiTokenAuthentication middleware) for third-party employer integrations at /api/v1/*'],
    ],
    [2.0, 4.5]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — AUTHENTICATION SYSTEM
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '2. Authentication System')

add_heading(2, '2.1 Auth Stack')
simple_table(
    ['Component', 'Detail'],
    [
        ['Auth Library',        'Laravel Fortify (core auth actions) + Laravel Sanctum (API tokens)'],
        ['Permission Layer',    'Spatie laravel-permission ^6.22 (roles & abilities)'],
        ['Two-Factor Auth',     'Laravel Fortify TwoFactorAuthenticatable trait + custom UI at /two-factor-authentication'],
        ['Social Auth',         'Laravel Socialite 5.x — Google, LinkedIn, Apple, Microsoft, Facebook, GitHub, Twitter'],
        ['Session Type',        'Database-backed PHP session (SESSION_DRIVER=database) for web; Sanctum personal access tokens for API'],
        ['Web Session Fields',  'user.id, user.name, user.email, user.account_type, user.is_active, user.company_id, user.preferences (JSON)'],
        ['Token Fields',        'Personal access tokens stored in personal_access_tokens table with tokenable_type, tokenable_id, name, token (hashed), abilities, expires_at'],
        ['Login Logic Path',    'app/Http/Controllers/Auth/AuthenticatedSessionController.php'],
        ['Registration Path',   'app/Http/Controllers/Auth/RegisteredUserController.php'],
        ['Rate Limiting',       'Login: 5 attempts/min; Registration: 3 attempts/min; Password reset: 5 attempts/hour'],
    ],
    [2.2, 4.3]
)

add_heading(2, '2.2 Organisation / Company Concept')
add_para('Yes — a Company model exists and is central to the employer side of the platform.')
add_bullet('Users have an account_type column: job_seeker | employer | admin')
add_bullet('Employer users have a company_id FK linking them to the companies table')
add_bullet('Each employer user belongs to exactly ONE company (no multi-seat / multi-admin per company yet)')
add_bullet('The Company model (app/Models/Company.php) stores: name, slug, logo, website, industry, company_size, founded_year, headquarters, description, is_verified, is_featured, ratings, review stats, contact details, LinkedIn URL, etc.')
add_bullet('No tenant isolation — all data is global; access is filtered by company_id in queries')
add_bullet('No concept of workspace, organisation, or true multi-tenancy exists')
add_warning('Only ONE user can be the employer for a given company (no team/seats model). Multiple admins per company is NOT supported.')

add_heading(2, '2.3 Users Table Columns')
add_para('From migrations (0001_01_01_000000_create_users_table.php + subsequent alter migrations):')
cols = [
    'id (bigint, PK, auto-increment)',
    'name (varchar 255)',
    'email (varchar 255, unique)',
    'email_verified_at (timestamp, nullable)',
    'password (varchar 255, bcrypt hashed)',
    'phone (varchar, nullable, encrypted at rest)',
    'account_type (enum: job_seeker / employer / admin, default: job_seeker)',
    'company_id (bigint, nullable, FK → companies.id)',
    'avatar (varchar, nullable)',
    'is_active (boolean, default: true)',
    'preferences (json, nullable)',
    'last_login_at (timestamp, nullable)',
    'timezone (varchar, default: UTC)',
    'two_factor_secret (text, nullable, encrypted)',
    'two_factor_recovery_codes (text, nullable, encrypted)',
    'two_factor_confirmed_at (timestamp, nullable)',
    'stripe_customer_id (varchar, nullable)',
    'remember_token (varchar 100, nullable)',
    'deleted_at (timestamp, nullable — soft deletes)',
    'created_at / updated_at (timestamps)',
]
for c in cols:
    add_bullet(c)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — DATABASE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '3. Database')

simple_table(
    ['Attribute', 'Value'],
    [
        ['ORM',                 'Eloquent ORM (Laravel)'],
        ['DB Type (Production)', 'MySQL (Azure Database for MySQL / managed)'],
        ['DB Type (.env.example default)', 'SQLite (development default — overridden in production)'],
        ['DB Version',          'Not pinned in codebase; Azure managed MySQL typically 8.x'],
        ['Migration Count',     '~137 migrations (as of 2026-05-29 deployment)'],
        ['Schema Prefix',       'None — all tables in single database, no schema namespacing'],
    ],
    [2.5, 4.0]
)

add_heading(2, '3.1 Complete Table / Model Inventory')
add_para('All models live in app/Models/. The corresponding database table names are listed below (grouped by domain):')

domains = {
    'Core Auth & Users': [
        'users', 'profiles', 'social_accounts', 'social_auth_logs', 'social_providers',
        'two_factor_authentications', 'api_tokens', 'personal_access_tokens',
    ],
    'Jobs & Applications': [
        'job_listings (Model: Job)', 'applications', 'saved_jobs', 'job_alerts',
        'job_views', 'job_embeddings', 'job_sources', 'job_templates',
        'discovered_jobs', 'job_matches', 'auto_applications',
    ],
    'Companies': [
        'companies', 'company_dna_profiles', 'company_intelligence_profiles',
        'company_interview_data', 'company_reviews', 'company_review_reports',
        'company_review_votes', 'company_team_members', 'company_blacklists',
        'employer_brand_scores', 'salary_reports', 'interview_experiences',
        'interview_experience_votes',
    ],
    'Subscriptions & Payments': [
        'subscription_plans', 'user_subscriptions', 'payment_transactions',
        'payment_activity_logs',
    ],
    'AI / Negotiation': [
        'ai_conversations', 'ai_usage_logs (ai_credit_logs)', 'ai_credit_logs',
        'ai_prompts', 'ai_decision_logs', 'ai_golden_tests', 'ai_golden_test_runs',
        'ai_resume_generations', 'negotiation_strategies', 'negotiation_sessions',
        'negotiation_messages', 'negotiation_scenarios', 'negotiation_scripts',
        'negotiation_tactics', 'counter_offers', 'offer_comparisons',
        'offer_letters', 'offer_letter_activities', 'offer_letter_templates',
        'human_overrides', 'ai_bias_reports', 'ai_disclaimers',
        'ai_disclaimer_acknowledgments',
    ],
    'Resume Builder': [
        'resumes', 'resume_templates', 'resume_versions', 'resume_ai_suggestions',
        'resume_analytics', 'cover_letters',
    ],
    'Skill & Learning': [
        'user_skills', 'skill_assessments', 'skill_badges', 'skill_validations',
        'skill_gaps', 'learning_paths', 'learning_resources', 'learning_progress',
        'daily_challenges', 'user_daily_challenges', 'assessments',
        'assessment_questions', 'assessment_attempts', 'assessment_responses',
        'assessment_refinements', 'vantage_skill_awards', 'coaching_skill_scores',
        'certificates',
    ],
    'Interview Intelligence': [
        'interviews', 'interview_sessions', 'interview_questions', 'interview_responses',
        'interview_feedbacks', 'interview_requests', 'interview_panel_scores',
        'interview_performance_reports', 'interview_coaching_tips', 'video_interview_sessions',
        'video_interview_recordings', 'video_interview_rooms', 'video_interview_templates',
        'video_interview_questions', 'video_interview_participants',
        'video_interview_invitations', 'video_interview_analyses', 'hiring_rounds',
        'hiring_tests', 'hiring_test_attempts', 'round_attempts',
        'evaluation_sessions', 'evaluation_answers', 'situational_scenarios',
        'scenario_responses',
    ],
    'Autonomous Agent': [
        'agent_configurations', 'agent_audit_logs', 'agent_internal_matches',
        'agent_learning_metrics', 'idempotency_keys',
    ],
    'ATS Integration': [
        'ats_connections', 'ats_providers', 'ats_synced_jobs', 'ats_job_mappings',
        'ats_candidate_mappings', 'ats_sync_logs', 'ats_webhooks',
    ],
    'Networking & Social': [
        'user_posts', 'post_comments', 'post_likes', 'comment_likes', 'user_follows',
        'connections', 'conversations', 'conversation_participants', 'messages',
        'message_reads', 'network_events', 'network_conversations', 'network_messages',
        'network_notification_settings', 'groups', 'group_members', 'group_posts',
        'hashtags',
    ],
    'Gamification': [
        'gamification_activities', 'gamification_badges', 'gamification_referral_bonuses',
        'user_gamification_profiles', 'points_transactions', 'xp_transactions',
        'achievements', 'user_achievements', 'badges', 'user_badges', 'rewards',
        'user_rewards', 'leaderboards', 'leaderboard_entries', 'seasonal_events',
        'user_event_participations',
    ],
    'Talent Marketplace (Freelancer)': [
        'marketplace_projects', 'marketplace_proposals', 'marketplace_contracts',
        'marketplace_milestones', 'marketplace_escrows', 'marketplace_messages',
        'marketplace_reviews', 'marketplace_disputes', 'marketplace_invitations',
        'marketplace_time_logs', 'freelancer_profiles', 'freelancer_gigs',
        'saved_freelancers', 'saved_projects',
    ],
    'Career Coaching': [
        'career_coach_sessions', 'career_coach_messages', 'career_coach_preferences',
        'career_coach_checkins', 'career_coach_suggestions', 'career_goals',
        'career_goal_updates', 'career_path_nodes', 'career_path_edges',
    ],
    'Calendar': [
        'calendar_connections', 'calendar_sync_events', 'scheduled_events',
        'scheduling_links', 'event_rsvps', 'event_reminders', 'event_participants',
    ],
    'Email Templates': [
        'email_templates', 'email_template_categories', 'email_template_versions',
        'email_template_analytics', 'email_sends', 'email_ai_customizations',
        'bulk_email_logs', 'message_templates',
    ],
    'Mentorship': [
        'mentor_profiles', 'mentorships', 'mentorship_matches',
    ],
    'Analytics & Predictive': [
        'talent_pipelines', 'talent_pool_candidates', 'talent_need_predictions',
        'success_predictions', 'success_indicators', 'hire_performances',
        'hiring_patterns', 'time_to_hire_metrics', 'productivity_estimates',
        'tenure_forecasts', 'flight_risk_assessments', 'team_dynamics',
        'job_market_heatmaps', 'application_funnels', 'source_attributions',
        'anonymized_screenings', 'bias_audit_results', 'fairness_metrics',
        'proxy_discrimination_alerts', 'hiring_decision_overrides', 'decision_traces',
    ],
    'Background Checks': [
        'background_checks', 'background_check_items', 'background_check_packages',
        'background_check_activities', 'background_check_adverse_actions',
        'background_check_webhooks',
    ],
    'Misc / System': [
        'audit_logs', 'feature_flags', 'testimonials', 'newsletters',
        'notifications', 'sent_notifications', 'notification_preferences',
        'push_subscriptions', 'webhooks', 'webhook_deliveries',
        'passive_candidate_profiles', 'silver_medalists', 'user_availability',
        'behavioral_assessments', 'application_notes', 'application_activity_logs',
        'application_status_histories', 'application_templates',
        'pipeline_candidates', 'talent_need_predictions',
        'benefits_packages', 'culture_analyses',
        'telescope_entries', 'cache', 'sessions', 'jobs (Laravel queue)', 'migrations',
        'permissions', 'roles', 'model_has_permissions', 'model_has_roles', 'role_has_permissions',
    ],
}

for domain, tables in domains.items():
    add_para(domain, bold=True)
    for t in tables:
        add_bullet(t, indent_level=0)

add_heading(2, '3.2 Key Model Column Details')

add_para('User (app/Models/User.php)', bold=True)
add_note('Full column list in Section 2.3 above.')

add_para('Company (app/Models/Company.php)', bold=True)
company_cols = [
    'id, name, slug, logo, logo_url, website, industry, company_size, founded_year',
    'headquarters, description, is_verified (bool), is_featured (bool), rating (float)',
    'follower_count (int), review_count, salary_count, interview_count',
    'avg_overall_rating, avg_culture_rating, avg_compensation_rating, avg_worklife_rating',
    'avg_growth_rating, avg_management_rating (floats)',
    'ceo_approval_rate, recommend_rate (ints)',
    'benefits (json array)',
    'company_email, hr_email, contact_phone, linkedin_url, culture, deleted_at, created_at, updated_at',
]
for c in company_cols:
    add_bullet(c)

add_para('Job / job_listings (app/Models/Job.php)', bold=True)
job_cols = [
    'id, company_id (FK), posted_by (user_id FK), title, slug',
    'description, location, location_type, work_mode (remote/hybrid/onsite)',
    'employment_type (full-time/part-time/contract/internship/remote/freelance/temporary)',
    'experience_level (entry/mid/senior/lead/executive)',
    'salary_min, salary_max, salary_currency, salary_period',
    'required_skills (json), preferred_skills (json), requirements (json)',
    'responsibilities (json), qualifications (text), benefits (json)',
    'application_method, external_url, application_email, application_instructions',
    'status (draft/active/published/paused/closed/archived)',
    'is_featured (bool), is_urgent (bool)',
    'published_at, expires_at, filled_at (datetimes)',
    'views_count, applications_count, saves_count (ints)',
    'search_keywords, ai_embeddings (json)',
    '--- Orin™ fields ---',
    'application_link_token, open_date, close_date, eval_start_date, final_date',
    'target_hire_count (int), orin_generated_jd (json), orin_application_form_fields (json)',
    'application_phase, requires_portfolio, requires_github, requires_work_sample (bools)',
    'mandatory_screening_questions (json), nice_to_have (json)',
    'ai_insights (json), quality_score',
    'deleted_at, created_at, updated_at',
]
for c in job_cols:
    add_bullet(c)

add_para('Application (app/Models/Application.php)', bold=True)
app_cols = [
    'id, user_id (FK), job_id (FK), application_number',
    'cover_letter (text), resume_file (varchar), answers (json), screening_answers (json)',
    'status (submitted/pending/viewed/reviewing/shortlisted/hired/rejected/withdrawn)',
    'status_updated_at, status_history (json), rejection_reason, ai_reason',
    'match_analysis (json), timeline (json), notes (text)',
    'submitted_at, viewed_at, is_archived (bool), source',
    '--- Orin™ evaluation ---',
    'evaluation_status, evaluation_score, skill_match_score, resume_quality_score',
    'behavioural_fit_score, final_rank_score (decimals), rank_position (int)',
    'evaluation_started_at, evaluation_completed_at',
    'application_email_sent, evaluation_invite_sent, result_email_sent (bools)',
    'result_notified_at, portfolio_url, github_url, work_sample_url',
    'access_token, is_guest_applicant (bool), guest_name, guest_email, guest_phone',
    '--- Hiring pipeline ---',
    'hiring_stage, pipeline_stage_date, pipeline_stage_notes',
    'confirmation_email_sent, test_link_token',
    'deleted_at, created_at, updated_at',
]
for c in app_cols:
    add_bullet(c)

add_heading(2, '3.3 company_id / tenant_id Occurrences')
add_para('No tenant_id or org_id column exists anywhere. company_id is used throughout:')
simple_table(
    ['Table / Context', 'Column', 'File Path'],
    [
        ['users',                         'company_id (FK → companies)', 'migration: 2025_11_29_082159_add_company_id_to_users_table.php'],
        ['job_listings',                  'company_id (FK → companies)', 'migration: 2025_10_28_162800_create_job_listings_table.php'],
        ['company_industry (pivot)',       'company_id (FK → companies)', 'migration: 2025_10_28_162758_create_company_industry_table.php'],
        ['company_dna_profiles (Scout)',   'company_id (FK → companies)', 'migration: 2025_11_06_000002_create_scout_corporate_dna_tables.php'],
        ['behavioral_assessment_configs', 'company_id (FK → companies)', 'migration: 2025_11_06_000004_create_scout_behavioral_assessment_tables.php'],
        ['question_banks',                'company_id (FK → companies)', 'migration: 2025_11_06_000003_create_scout_assessment_tables.php'],
        ['Multiple employer portal tables','company_id (FK → companies)', 'migrations: 2026_01_16_000000, 2026_01_17_000000'],
        ['EmployerDashboardController',   'company_id (via user→company relationship)', 'app/Http/Controllers/Employer/EmployerDashboardController.php'],
        ['ApplicantTrackingController',   'company_id (JOIN job_listings)', 'app/Http/Controllers/Employer/ApplicantTrackingController.php'],
        ['JobPostingController',          'company_id (WHERE clause)', 'app/Http/Controllers/Employer/JobPostingController.php'],
    ],
    [2.2, 1.5, 2.8]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — AI & AZURE OPENAI INTEGRATION
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '4. AI and Azure OpenAI Integration')

add_heading(2, '4.1 AI Provider Configuration')
simple_table(
    ['Setting', 'Value'],
    [
        ['Primary Provider',      'Azure OpenAI (config key: ai.azure)'],
        ['Primary Model',         'gpt-5.4 (deployment ID: gpt-5.4)'],
        ['Primary Endpoint',      'https://studai-openai-2049701603.openai.azure.com/'],
        ['API Version',           '2025-04-01-preview'],
        ['Fallback Provider',     'Azure Anthropic'],
        ['Fallback Model',        'claude-sonnet-4-6'],
        ['Embedding Model',       'text-embedding-3-large'],
        ['Default Max Tokens',    '16,384'],
        ['Default Temperature',   '0.7'],
        ['Legacy OpenAI',         'openai-php/laravel ^0.17.1 installed but only used as compatibility shim via OpenAIService → AIService'],
    ],
    [2.2, 4.3]
)

add_warning('CRITICAL SECURITY: Azure OpenAI API key (e6e0bf0c61d14319a92bbc2d2a02f52f) and endpoint are HARDCODED as fallback values in config/ai.php lines 26-28 and in AIService.php lines 47-51. If AZURE_OPENAI_API_KEY env var is empty, the hardcoded key is used. This key must be rotated and removed from source code immediately.')
add_warning('Azure OpenAI endpoint URL is also hardcoded in the same location.')

add_heading(2, '4.2 Files Making AI API Calls')

ai_files = [
    ('app/Services/AI/AIService.php',                  'callAzureOpenAI(), callAnthropicAI(), generateText(), generateJSON()',                                       'gpt-5.4 / claude-sonnet-4-6', 'Central AI gateway — all other services route through this'),
    ('app/Services/AI/OpenAIService.php',              'generateCompletion()',                                                                                       'Delegates to AIService',      'Compatibility shim extending AIService'),
    ('app/Services/AI/NegotiationStrategistService.php','generateStrategy(), gatherMarketResearch(), analyzeNegotiationStrength(), getCompanyIntelligence(), generateAiInsights()', 'gpt-5.4', 'Full salary negotiation strategy generation'),
    ('app/Services/AI/NegotiationCoachingService.php', 'sendInitialGuidance(), respondToMessage(), generateScenario()',                                             'gpt-5.4',   'Live negotiation coaching conversation'),
    ('app/Services/AI/NegotiationScenarioService.php', 'generateScenario(), evaluateResponse()',                                                                    'gpt-5.4',   'Practice negotiation scenario generation'),
    ('app/Services/AI/NegotiationScriptService.php',   'generateScript()',                                                                                          'gpt-5.4',   'Word-for-word negotiation script generation'),
    ('app/Services/AI/CandidateScreeningService.php',  'analyzeCandidate(), analyzeSkillMatch(), analyzeCultureFit(), generateRecommendation(), generateInterviewQuestions()', 'gpt-5.4', 'AI-powered applicant screening & scoring'),
    ('app/Services/AI/ResumeAIService.php',            'generateSummary(), extractSkills(), optimizeForJob(), analyzeATS()',                                         'gpt-5.4',   'Resume AI analysis and optimisation'),
    ('app/Services/AI/ResumeAnalyzerService.php',      'analyzeResume(), scoreResume()',                                                                            'gpt-5.4',   'Deep resume quality analysis'),
    ('app/Services/AI/ResumeCustomizationService.php', 'customizeForJob()',                                                                                         'gpt-5.4',   'Tailors resume content to specific job'),
    ('app/Services/AI/CoverLetterGeneratorService.php','generateCoverLetter()',                                                                                     'gpt-5.4',   'AI cover letter generation'),
    ('app/Services/AI/InterviewPrepService.php',       'generatePrepPlan(), generateMockQuestions()',                                                               'gpt-5.4',   'Interview preparation content'),
    ('app/Services/AI/InterviewQuestionGenerator.php', 'generateQuestions(), generateFollowUp()',                                                                   'gpt-5.4',   'Dynamic interview question generation'),
    ('app/Services/AI/AnswerEvaluationService.php',    'evaluateAnswer(), generateFeedback()',                                                                      'gpt-5.4',   'Real-time interview answer scoring'),
    ('app/Services/AI/CareerCoachService.php',         'generateResponse(), suggestGoals(), analyseTrend()',                                                        'gpt-5.4',   'AI career coach conversation engine'),
    ('app/Services/AI/CareerAdvisorService.php',       'advise(), generateRoadmap()',                                                                               'gpt-5.4',   'Career path advisory'),
    ('app/Services/AI/CareerTrajectoryService.php',    'predictTrajectory()',                                                                                       'gpt-5.4',   'Career progression prediction'),
    ('app/Services/AI/JobMatchingService.php',         'matchJobs(), scoreMatch()',                                                                                 'gpt-5.4',   'Semantic job-to-candidate matching'),
    ('app/Services/AI/EmbeddingService.php',           'embed(), getEmbedding()',                                                                                   'text-embedding-3-large', 'Vector embeddings for semantic search'),
    ('app/Services/AI/SkillGapAnalyzerService.php',    'analyzeGaps(), prioritizeGaps()',                                                                           'gpt-5.4',   'Skill gap identification and learning path'),
    ('app/Services/AI/SkillAssessmentGeneratorService.php','generateAssessment()',                                                                                  'gpt-5.4',   'Dynamic skill assessment questions'),
    ('app/Services/AI/SkillValidatorService.php',      'validate()',                                                                                                'gpt-5.4',   'AI-powered skill claim validation'),
    ('app/Services/AI/SkillTrendAnalysisService.php',  'analyzeMarketTrends()',                                                                                     'gpt-5.4',   'Industry skill demand analysis'),
    ('app/Services/AI/SalaryIntelligenceService.php',  'getSalaryIntelligence(), benchmarkSalary()',                                                               'gpt-5.4',   'Market salary benchmarking'),
    ('app/Services/AI/OrinEvaluationService.php',      'evaluateApplication(), scoreCandidate()',                                                                   'gpt-5.4',   'Orin™ automated applicant evaluation pipeline'),
    ('app/Services/AI/OrinJobCreatorService.php',      'createJob(), generateJD()',                                                                                 'gpt-5.4',   'AI-powered job description generation'),
    ('app/Services/AI/CompanyResearchService.php',     'researchCompany()',                                                                                         'gpt-5.4',   'Company intelligence gathering'),
    ('app/Services/AI/ApplicationOptimizerService.php','optimize()',                                                                                                'gpt-5.4',   'Application quality optimisation'),
    ('app/Services/AI/LearningPathCuratorService.php', 'curatePath()',                                                                                              'gpt-5.4',   'Personalised learning path curation'),
    ('app/Services/AI/SpeechToTextService.php',        'transcribe()',                                                                                              'Azure STT / Whisper', 'Video interview transcription'),
    ('app/Services/AI/PromptRegistryService.php',      'get(), render(), getSystemPrompt(), getConfig()',                                                           'N/A (meta)',  'Manages DB-stored prompts with versioning and A/B testing'),
    ('app/Services/AI/Scout/ (multiple files)',         'Various Scout™ assessment and DNA analysis methods',                                                        'gpt-5.4',   'Corporate DNA, behavioural assessments, passive candidate discovery'),
]

simple_table(
    ['File Path', 'Key Method(s)', 'Model Called', 'Purpose'],
    [[a,b,c,d] for a,b,c,d in ai_files],
    [2.5, 1.8, 1.2, 1.6]
)

add_heading(2, '4.3 System Prompts')
add_para('Prompts are managed via two mechanisms:')
add_bullet('DB-stored prompts: AIPrompt model (ai_prompts table) managed via PromptRegistryService. Prompts have name, system_prompt, template, variables, version, model_hint, temperature, max_tokens, is_active columns. Can be edited without code deployment.')
add_bullet('Inline prompts: Embedded as PHP strings directly in service methods (e.g., NegotiationStrategistService::gatherMarketResearch uses a hardcoded prompt string for salary benchmarking). These require code changes to modify.')
add_note('Example inline system prompt (NegotiationStrategistService): "You are a compensation intelligence analyst with deep knowledge of the Indian job market (2024-2025)..."')
add_note('All AI calls pass through AIService which builds the messages array. System prompts are passed as role: "system" messages. User input (resume text, job data, salary figures) is passed as role: "user" messages.')

add_heading(2, '4.4 AI Environment Variables')
simple_table(
    ['Variable Name', 'Purpose'],
    [
        ['AI_PRIMARY_PROVIDER',         'Primary AI provider: azure or anthropic'],
        ['AI_FALLBACK_PROVIDER',        'Fallback AI provider'],
        ['AZURE_OPENAI_API_KEY',        'Azure OpenAI API key (primary)'],
        ['AZURE_OPENAI_ENDPOINT',       'Azure OpenAI resource endpoint URL'],
        ['AZURE_OPENAI_DEPLOYMENT_ID',  'Azure deployment name (model)'],
        ['AZURE_OPENAI_API_VERSION',    'Azure OpenAI API version string'],
        ['AZURE_OPENAI_MODEL',          'Chat model name (e.g., gpt-5.4)'],
        ['AZURE_OPENAI_MODEL_MINI',     'Lighter chat model for cheaper tasks'],
        ['AI_MODEL_EMBEDDINGS',         'Embeddings model (text-embedding-3-large)'],
        ['AZURE_ANTHROPIC_API_KEY',     'Azure Anthropic API key (fallback)'],
        ['AZURE_ANTHROPIC_ENDPOINT',    'Azure Anthropic endpoint URL'],
        ['AZURE_ANTHROPIC_MODEL',       'Anthropic model (claude-sonnet-4-6)'],
        ['ANTHROPIC_MAX_TOKENS',        'Max tokens for Anthropic calls'],
        ['AI_MAX_TOKENS',               'Default max tokens for AI requests'],
        ['AI_TEMPERATURE',              'Default temperature (0.0–1.0)'],
        ['AI_CACHE_ENABLED',            'Enable/disable AI response caching'],
        ['AI_REQUEST_TIMEOUT',          'Default request timeout (seconds)'],
        ['AI_EMBEDDINGS_TIMEOUT',       'Embeddings request timeout'],
        ['AI_STREAMING_TIMEOUT',        'Streaming response timeout'],
        ['AI_LONG_RUNNING_TIMEOUT',     'Long-running AI task timeout'],
        ['OPENAI_API_KEY',              'Legacy OpenAI direct key (optional)'],
        ['OPENAI_ORGANIZATION',         'Legacy OpenAI org ID (optional)'],
    ],
    [2.5, 4.0]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — API ROUTES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '5. API Routes')

add_heading(2, '5.1 Public Routes (No Authentication)')
simple_table(
    ['Method', 'Path', 'Controller@Method', 'Description'],
    [
        ['GET', '/api/health',                     'Closure',                          'Health check endpoint — returns {status, version, timestamp}'],
        ['GET', '/api/skills/certificate/{hash}',  'SkillAnalyzerController@getCertificate', 'Public certificate verification by hash'],
        ['GET', '/',                               'MarketingController@home',          'Landing page'],
        ['GET', '/features, /about, /pricing, /for-employers, /contact', 'MarketingController@*', 'Marketing pages'],
        ['GET', '/jobs',                           'JobController@index',               'Public job search listing'],
        ['GET', '/jobs/{job}',                     'JobController@show',               'Public job detail page'],
        ['GET', '/apply/{token}',                  'PublicApplyController@show',       'Public job application form (tokenised)'],
        ['POST','/apply/{token}',                  'PublicApplyController@store',      'Submit public application'],
        ['GET', '/company/{slug}',                 'CompanyController@show',           'Public company profile'],
        ['GET', '/r/{shareToken}',                 'ResumeController@publicView',      'Public resume view via share token'],
    ],
    [0.7, 2.0, 2.5, 2.0]
)

add_heading(2, '5.2 Authenticated Routes (auth:sanctum — API)')
simple_table(
    ['Method', 'Path', 'Controller@Method', 'Description'],
    [
        ['GET',    '/api/user',                            'UserController@me',                    'Get current authenticated user'],
        ['GET',    '/api/jobs/recommended',                'JobMatchingController@recommended',    'AI-matched job recommendations'],
        ['GET',    '/api/jobs/search',                     'JobMatchingController@search',         'Keyword + semantic job search'],
        ['GET',    '/api/jobs/{job}/match-analysis',       'JobMatchingController@matchAnalysis',  'Detailed AI match score for a job'],
        ['GET',    '/api/jobs/saved',                      'JobMatchingController@saved',          'List saved jobs'],
        ['POST',   '/api/jobs/{job}/save',                 'JobMatchingController@save',           'Save a job'],
        ['DELETE', '/api/jobs/{job}/unsave',               'JobMatchingController@unsave',         'Unsave a job'],
        ['POST',   '/api/jobs/{job}/ai-apply',             'JobMatchingController@apply',          'One-click AI-powered application'],
        ['POST',   '/api/payment/initiate',                'PaymentController@initiate',           'Create payment order (idempotent)'],
        ['POST',   '/api/payment/razorpay/callback',       'PaymentController@razorpayCallback',   'Razorpay payment verification callback'],
        ['GET',    '/api/payment/history',                 'PaymentController@history',            'Payment transaction history'],
        ['POST',   '/api/payment/refund/{transaction}',    'PaymentController@requestRefund',      'Request refund (idempotent)'],
        ['GET/POST/PUT/DELETE', '/api/profile/*',          'ProfileController@*',                  'Profile CRUD + sections (experience, education, skills)'],
        ['POST',   '/api/interview/sessions',              'InterviewSessionController@start',     'Start AI interview session'],
        ['GET',    '/api/interview/sessions/{id}',         'InterviewSessionController@show',      'Get session details'],
        ['GET',    '/api/interview/sessions/{id}/next-question', 'InterviewSessionController@getNextQuestion', 'Get AI-generated next question'],
        ['POST',   '/api/interview/sessions/{id}/answer', 'InterviewSessionController@submitAnswer', 'Submit answer for AI evaluation'],
        ['GET',    '/api/interview/sessions/{id}/report', 'InterviewSessionController@getReport', 'Get full interview performance report'],
        ['GET/POST','/api/agent/*',                        'AgentController@*',                    'Autonomous agent config, activate, pause, status, metrics'],
        ['GET/POST','/api/negotiation/*',                  'NegotiationController@*',              'Negotiation strategy creation and coaching'],
        ['POST',   '/api/skills/analyze',                  'SkillAnalyzerController@analyzeSkillGaps', 'Trigger skill gap analysis'],
        ['GET',    '/api/skills/gaps',                     'SkillAnalyzerController@listSkillGaps','List user skill gaps'],
        ['POST',   '/api/skills/learning-path/{gapId}',   'SkillAnalyzerController@generateLearningPath', 'Generate learning path for skill gap'],
        ['POST',   '/api/skills/assessment/{skillId}',    'SkillAnalyzerController@generateAssessment', 'Generate skill assessment'],
        ['GET/POST','/api/gdpr/*',                         'GDPRController@*',                     'GDPR data export, deletion requests'],
    ],
    [0.8, 2.2, 2.0, 2.0]
)

add_heading(2, '5.3 Third-Party Employer API (Custom Token Auth)')
add_para('Prefix: /api/v1/ — uses ApiTokenAuthentication + ApiRateLimiting middleware. Ability-based access control via ApiAbilityCheck.')
simple_table(
    ['Method', 'Path', 'Abilities Required', 'Description'],
    [
        ['GET',    '/api/v1/company',                    'company.read',        'Get employer company profile'],
        ['GET',    '/api/v1/company/statistics',         'company.read',        'Company hiring statistics'],
        ['PUT',    '/api/v1/company',                    'company.write',       'Update company profile'],
        ['GET',    '/api/v1/jobs',                       'jobs.read',           'List company job postings'],
        ['GET',    '/api/v1/jobs/{job}',                 'jobs.read',           'Get job detail'],
        ['GET',    '/api/v1/jobs/{job}/statistics',      'jobs.read',           'Job application statistics'],
        ['POST',   '/api/v1/jobs',                       'jobs.write',          'Create job posting'],
        ['PUT',    '/api/v1/jobs/{job}',                 'jobs.write',          'Update job posting'],
        ['DELETE', '/api/v1/jobs/{job}',                 'jobs.write',          'Delete job posting'],
        ['GET',    '/api/v1/applications',               'applications.read',   'List applications'],
        ['GET',    '/api/v1/applications/{app}',         'applications.read',   'Application detail'],
        ['PUT',    '/api/v1/applications/{app}/status',  'applications.write',  'Update application status'],
        ['POST',   '/api/v1/applications/bulk-status',   'applications.write',  'Bulk status update'],
    ],
    [0.7, 2.0, 1.5, 2.3]
)

add_heading(2, '5.4 Admin-Only Routes')
add_para('Prefix: /admin/ — middleware: auth:admin (EnsureUserIsAdmin)')
simple_table(
    ['Method', 'Path', 'Controller@Method', 'Description'],
    [
        ['GET', '/admin/analytics',               'AdminAnalyticsController@dashboard',     'Admin analytics overview'],
        ['GET', '/admin/analytics/revenue',       'AdminAnalyticsController@revenue',       'Revenue analytics'],
        ['GET', '/admin/analytics/subscriptions', 'AdminAnalyticsController@subscriptions', 'Subscription analytics'],
        ['GET', '/admin/analytics/users',         'AdminAnalyticsController@users',         'User growth analytics'],
        ['GET', '/admin/analytics/applications',  'AdminAnalyticsController@applications',  'Application analytics'],
        ['GET', '/admin/analytics/churn',         'AdminAnalyticsController@churn',         'Churn analytics'],
        ['GET', '/admin/analytics/export',        'AdminAnalyticsController@export',        'Export analytics data'],
        ['GET', '/studai/*',                      'Filament Admin Panel',                   'Full Filament admin (admin account_type only)'],
    ],
    [0.7, 2.0, 2.2, 2.1]
)

add_heading(2, '5.5 Webhook Routes (No Auth — Signature Verified)')
simple_table(
    ['Method', 'Path', 'Controller@Method', 'Description'],
    [
        ['POST', '/webhooks/stripe',    'StripeWebhookController@handleWebhook',   'Stripe events: checkout.session.completed, payment_intent.succeeded, charge.refunded'],
        ['POST', '/webhooks/razorpay', 'PaymentWebhookController@razorpayWebhook', 'Razorpay events: payment.captured, payment.failed'],
        ['GET/POST', '/ats/webhook/*', 'AtsController@*',                          'ATS provider integration webhooks'],
    ],
    [0.7, 1.8, 2.2, 2.3]
)

add_heading(2, '5.6 API Versioning')
add_para('Partial versioning: /api/v1/ prefix exists exclusively for the third-party employer integration API. All internal Sanctum routes (/api/*) are unversioned. No /api/v2/ exists.')

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — KEY BUSINESS LOGIC
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '6. Key Business Logic')

add_heading(2, '6.1 Main Entity')
add_para('The platform manages two core entities with equal importance:')
add_bullet('Job Listings (job_listings table / Job model) — the central matching surface')
add_bullet('Applications (applications table) — the core transactional entity connecting candidates to jobs')
add_para('Supporting entities drive the AI features: Resume, NegotiationStrategy, InterviewSession, AgentConfiguration.')

add_heading(2, '6.2 Five Most Important Controllers')
simple_table(
    ['Controller', 'Path', 'What It Does'],
    [
        ['EmployerDashboardController', 'app/Http/Controllers/Employer/EmployerDashboardController.php',
         'Employer home page: job counts, application statistics, recent applicants, hiring activity, intelligence profile. Uses Redis caching (300-600s TTL) and single GROUP BY queries.'],
        ['ApplicantTrackingController', 'app/Http/Controllers/Employer/ApplicantTrackingController.php',
         'Full ATS: list/filter applicants, update status (shortlist/hire/reject), bulk actions, schedule interviews, compare candidates, add notes, export CSV. Triggers notifications and queued hiring emails.'],
        ['JobPostingController', 'app/Http/Controllers/Employer/JobPostingController.php',
         'Employer job CRUD: create/edit/publish/close job listings. Includes AI-content generation (description, requirements, questions from just the job title). Manages Orin™ evaluation pipeline dates.'],
        ['DashboardController', 'app/Http/Controllers/DashboardController.php',
         'Job seeker home page: application stats, recent applications, recommended jobs, subscription status, profile completion. Redirects employers to employer.dashboard.'],
        ['ResumeController', 'app/Http/Controllers/ResumeController.php',
         'Resume CRUD + AI features: generate summary, extract skills, ATS optimise, export PDF/DOCX. Uses ResumeAIService, ResumeExportService. Authorises ownership via Policy.'],
        ['AgentController (API)', 'app/Http/Controllers/API/AgentController.php',
         'Autonomous job agent: configure preferences, activate/pause/resume/deactivate, monitor status, view auto-applications, manage company blacklist, trigger manual discovery.'],
        ['NegotiationController (API)', 'app/Http/Controllers/API/NegotiationController.php',
         'Salary negotiation: create strategy, start live coaching session, send messages, practice scenarios, generate word-for-word scripts, compare offers.'],
    ],
    [1.8, 2.0, 3.2]
)

add_heading(2, '6.3 Queue Jobs')
simple_table(
    ['Job Class', 'Path', 'What It Does'],
    [
        ['DiscoverJobsJob',                 'app/Jobs/Agent/',           'Fetches jobs from RSS/API feeds (RemoteOK, WeWorkRemotely, HackerNews, Jobicy etc.) for autonomous agent users'],
        ['ScanInternalJobsJob',             'app/Jobs/Agent/',           'Scans internal platform job listings for agent matches'],
        ['SubmitApplicationsJob',           'app/Jobs/Agent/',           'Autonomously submits applications on behalf of users (requires human approval if configured)'],
        ['FollowUpJob',                     'app/Jobs/Agent/',           'Sends follow-up messages for pending applications'],
        ['SendDigestJob',                   'app/Jobs/Agent/',           'Sends daily/weekly digest of agent activity to users'],
        ['UpdateLearningJob',               'app/Jobs/Agent/',           'Updates ML learning metrics from agent outcomes'],
        ['AutomatedShortlistingJob',        'app/Jobs/',                 'AI-powered bulk candidate shortlisting for a job posting'],
        ['GenerateCandidateMatchScoresJob', 'app/Jobs/',                 'Computes AI match scores for all candidates on a job'],
        ['ScoreAndRankCandidates',          'app/Jobs/',                 'Final ranking pass across all evaluated candidates'],
        ['GenerateJobEmbeddings',           'app/Jobs/',                 'Generates text-embedding-3-large vectors for a job (stored in job_embeddings)'],
        ['SendHiringEmailsJob',             'app/Jobs/',                 'Sends AI-personalised hiring outcome emails (shortlisted/hired/rejected)'],
        ['SendShortlistNotifications',      'app/Jobs/',                 'Sends in-app + email notifications to shortlisted candidates'],
        ['AnalyzeSkillGapsJob',             'app/Jobs/',                 'Async skill gap analysis for a user'],
        ['CurateLearningResourcesJob',      'app/Jobs/',                 'Curates personalised learning resources for skill gaps'],
        ['GenerateAssessmentJob',           'app/Jobs/',                 'Generates AI skill assessment questions'],
        ['GenerateBehavioralAssessmentJob', 'app/Jobs/',                 'Generates behavioural assessment for Scout™'],
        ['AnalyzeCompanyDNAJob',            'app/Jobs/',                 'Analyses company culture/DNA from reviews and data'],
        ['AnalyzeVideoInterview',           'app/Jobs/',                 'Analyses recorded video interview (transcription + scoring)'],
        ['TranscribeVideoRecording',        'app/Jobs/',                 'Azure STT transcription of video interview recording'],
        ['ProcessVideoRecording',           'app/Jobs/',                 'Post-processing pipeline for video recordings'],
        ['TrackEmployerBrandJob',           'app/Jobs/',                 'Computes employer brand score from reviews, response rates, etc.'],
        ['UpdatePredictionsJob',            'app/Jobs/',                 'Refreshes predictive analytics (tenure forecast, flight risk, success predictions)'],
        ['UpdateTalentPipelinesJob',        'app/Jobs/',                 'Updates talent pipeline stage assignments'],
        ['DiscoverPassiveCandidatesJob',    'app/Jobs/',                 'AI-powered passive candidate identification'],
        ['RetryFailedPaymentJob',           'app/Jobs/',                 'Retries failed subscription payment with backoff'],
        ['ProcessJobAlerts',               'app/Jobs/',                 'Sends personalised job alert emails to subscribed users'],
        ['SendBulkEmail',                   'app/Jobs/',                 'Sends bulk employer outreach emails to talent pool'],
        ['GenerateUserDataExport',          'app/Jobs/',                 'GDPR data export generation'],
        ['AuditBiasJob',                    'app/Jobs/',                 'Runs bias detection analysis on AI decisions'],
        ['MonitorModelDriftJob',            'app/Jobs/',                 'Monitors AI model performance drift'],
    ],
    [2.0, 1.0, 4.0]
)

add_heading(2, '6.4 Webhook Handlers')
simple_table(
    ['Handler', 'Path', 'Events Handled'],
    [
        ['StripeWebhookController',  'app/Http/Controllers/StripeWebhookController.php',  'checkout.session.completed, payment_intent.succeeded, payment_intent.payment_failed, charge.refunded. Verifies Stripe-Signature header.'],
        ['PaymentWebhookController', 'app/Http/Controllers/PaymentWebhookController.php', 'Razorpay: payment.captured, payment.failed. Verifies HMAC-SHA256 signature against RAZORPAY_WEBHOOK_SECRET.'],
        ['AtsController',            'app/Http/Controllers/AtsController.php',            'ATS provider webhooks (Greenhouse, Lever etc.) — candidate status sync, job sync events.'],
        ['BackgroundCheckWebhook',   'Via BackgroundCheckService',                        'Third-party background check provider callbacks (candidate report ready, adverse action notices).'],
    ],
    [1.8, 2.2, 3.0]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — ENVIRONMENT VARIABLES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '7. Environment Variables')

add_heading(2, '7.1 Full .env.example Variable List')
env_vars = {
    'Application': ['APP_NAME', 'APP_ENV', 'APP_KEY', 'APP_DEBUG', 'APP_URL', 'APP_LOCALE', 'APP_FALLBACK_LOCALE', 'APP_FAKER_LOCALE', 'APP_MAINTENANCE_DRIVER', 'PHP_CLI_SERVER_WORKERS', 'BCRYPT_ROUNDS'],
    'Logging': ['LOG_CHANNEL', 'LOG_STACK', 'LOG_DEPRECATIONS_CHANNEL', 'LOG_LEVEL'],
    'Database': ['DB_CONNECTION', 'DB_HOST', 'DB_PORT', 'DB_DATABASE', 'DB_USERNAME', 'DB_PASSWORD'],
    'Session / Cache / Queue': ['SESSION_DRIVER', 'SESSION_LIFETIME', 'SESSION_ENCRYPT', 'SESSION_PATH', 'SESSION_DOMAIN', 'BROADCAST_CONNECTION', 'FILESYSTEM_DISK', 'QUEUE_CONNECTION', 'CACHE_STORE', 'CACHE_PREFIX'],
    'Redis': ['REDIS_CLIENT', 'REDIS_HOST', 'REDIS_PASSWORD', 'REDIS_PORT'],
    'Email': ['MAIL_MAILER', 'MAIL_SCHEME', 'MAIL_HOST', 'MAIL_PORT', 'MAIL_USERNAME', 'MAIL_PASSWORD', 'MAIL_FROM_ADDRESS', 'MAIL_FROM_NAME'],
    'Storage (AWS / Azure)': ['AWS_ACCESS_KEY_ID', 'AWS_SECRET_ACCESS_KEY', 'AWS_DEFAULT_REGION', 'AWS_BUCKET', 'AWS_USE_PATH_STYLE_ENDPOINT', 'AZURE_STORAGE_ACCOUNT', 'AZURE_STORAGE_KEY', 'AZURE_STORAGE_CONTAINER'],
    'AI / OpenAI': ['AI_PRIMARY_PROVIDER', 'AI_FALLBACK_PROVIDER', 'AZURE_OPENAI_API_KEY', 'AZURE_OPENAI_ENDPOINT', 'AZURE_OPENAI_DEPLOYMENT_ID', 'AZURE_OPENAI_API_VERSION', 'AZURE_OPENAI_MODEL', 'AZURE_OPENAI_MODEL_MINI', 'AI_MODEL_EMBEDDINGS', 'AZURE_ANTHROPIC_API_KEY', 'AZURE_ANTHROPIC_ENDPOINT', 'AZURE_ANTHROPIC_MODEL', 'ANTHROPIC_MAX_TOKENS', 'AI_MAX_TOKENS', 'AI_TEMPERATURE', 'AI_CACHE_ENABLED', 'AI_REQUEST_TIMEOUT', 'AI_EMBEDDINGS_TIMEOUT', 'AI_STREAMING_TIMEOUT', 'AI_LONG_RUNNING_TIMEOUT', 'OPENAI_API_KEY', 'OPENAI_ORGANIZATION', 'OPENAI_REQUEST_TIMEOUT'],
    'Payments': ['PAYMENT_GATEWAY', 'RAZORPAY_KEY', 'RAZORPAY_SECRET', 'RAZORPAY_WEBHOOK_SECRET', 'RAZORPAY_CURRENCY', 'PAYU_MERCHANT_KEY', 'PAYU_MERCHANT_SALT', 'PAYU_MODE', 'PAYU_CURRENCY', 'PAYU_SURL', 'PAYU_FURL', 'STRIPE_KEY', 'STRIPE_SECRET', 'STRIPE_WEBHOOK_SECRET'],
    'Search': ['SCOUT_DRIVER', 'MEILISEARCH_HOST', 'MEILISEARCH_KEY'],
    'Monitoring / Queues': ['SENTRY_LARAVEL_DSN', 'SENTRY_ENVIRONMENT', 'SENTRY_TRACES_SAMPLE_RATE', 'SENTRY_PROFILES_SAMPLE_RATE', 'HORIZON_DOMAIN', 'HORIZON_PATH', 'HORIZON_PREFIX'],
    'Real-time (Reverb)': ['REVERB_APP_ID', 'REVERB_APP_KEY', 'REVERB_APP_SECRET', 'REVERB_HOST', 'REVERB_PORT', 'REVERB_SCHEME'],
    'Push Notifications': ['VAPID_PUBLIC_KEY', 'VAPID_PRIVATE_KEY'],
    'Social Auth (OAuth)': ['GOOGLE_CLIENT_ID', 'GOOGLE_CLIENT_SECRET', 'GOOGLE_REDIRECT', 'LINKEDIN_CLIENT_ID', 'LINKEDIN_CLIENT_SECRET', 'GITHUB_CLIENT_ID', 'GITHUB_CLIENT_SECRET', 'MICROSOFT_CLIENT_ID', 'MICROSOFT_CLIENT_SECRET'],
    'Auth / Security': ['AUTH_DIAG_TOKEN', 'SANCTUM_STATEFUL_DOMAINS'],
    'Video Interviews': ['VIDEO_INTERVIEW_PROVIDER', 'TWILIO_ACCOUNT_SID', 'TWILIO_AUTH_TOKEN', 'TWILIO_API_KEY', 'TWILIO_API_SECRET'],
    'Background Checks': ['BACKGROUND_CHECK_PROVIDER', 'BACKGROUND_CHECK_API_KEY', 'BACKGROUND_CHECK_WEBHOOK_SECRET'],
}

for group, vars_list in env_vars.items():
    add_para(group, bold=True)
    add_bullet(', '.join(vars_list))

add_heading(2, '7.2 Variables by Category')
simple_table(
    ['Category', 'Variable Names'],
    [
        ['AI / OpenAI',    'AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_DEPLOYMENT_ID, AZURE_OPENAI_API_VERSION, AZURE_OPENAI_MODEL, AZURE_OPENAI_MODEL_MINI, AI_MODEL_EMBEDDINGS, AZURE_ANTHROPIC_API_KEY, AZURE_ANTHROPIC_ENDPOINT, AZURE_ANTHROPIC_MODEL, ANTHROPIC_MAX_TOKENS, AI_MAX_TOKENS, AI_TEMPERATURE, AI_REQUEST_TIMEOUT, AI_PRIMARY_PROVIDER, AI_FALLBACK_PROVIDER, OPENAI_API_KEY (legacy)'],
        ['Payments',       'RAZORPAY_KEY, RAZORPAY_SECRET, RAZORPAY_WEBHOOK_SECRET, RAZORPAY_CURRENCY, PAYU_MERCHANT_KEY, PAYU_MERCHANT_SALT, PAYU_MODE, STRIPE_KEY, STRIPE_SECRET, STRIPE_WEBHOOK_SECRET, PAYMENT_GATEWAY'],
        ['Email',          'MAIL_MAILER, MAIL_HOST, MAIL_PORT, MAIL_USERNAME, MAIL_PASSWORD, MAIL_FROM_ADDRESS, MAIL_FROM_NAME, MAIL_SCHEME'],
        ['External APIs',  'GOOGLE_CLIENT_ID/SECRET, LINKEDIN_CLIENT_ID/SECRET, GITHUB_CLIENT_ID/SECRET, MICROSOFT_CLIENT_ID/SECRET, TWILIO_ACCOUNT_SID/AUTH_TOKEN, MEILISEARCH_HOST/KEY, SENTRY_LARAVEL_DSN, VAPID_PUBLIC_KEY/PRIVATE_KEY, REVERB_APP_KEY/SECRET, BACKGROUND_CHECK_API_KEY'],
    ],
    [1.5, 5.0]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — LIMITATIONS / GAPS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '8. Current Limitations / Obvious Gaps')

add_heading(2, '8.1 Multi-User / Multi-Tenancy')
add_warning('No multi-tenancy: The platform has NO workspace/org/tenant isolation. All employer users share one database with filtering by company_id only.')
add_warning('Single employer per company: Only one user account can be linked to each company (company_id FK on users). There is no team-seats model, no invitation system for colleagues, and no role differentiation within a company (e.g., HR vs. Hiring Manager vs. Admin).')
add_bullet('Job seekers: each account is independent — no multi-user consideration needed.')
add_bullet('Employers: a company exists as a shared entity but only one user (the employer) is linked to it.')
add_note('Adding multi-seat employer teams would require: a company_users pivot table, role-within-company column, invitations system, and scoping all employer queries through team membership.')

add_heading(2, '8.2 Hardcoded Values')
add_warning('CRITICAL: Azure OpenAI API key (e6e0bf0c61d14319a92bbc2d2a02f52f) hardcoded in config/ai.php (line 26) and AIService.php (line 51) as fallback values. Must be rotated and removed.')
add_warning('CRITICAL: Azure OpenAI endpoint URL (https://studai-openai-2049701603.openai.azure.com/) hardcoded in same locations.')
add_bullet('Razorpay theme colour hardcoded as "#ec4899" in config/payment.php — should be in config.')
add_bullet('Currency defaults to "INR" throughout — no internationalisation support for non-Indian markets.')
add_bullet('Salary benchmarks in NegotiationStrategistService prompt are Indian market specific and hardcoded as prompt text.')
add_bullet('Auth diagnostic route (/auth-diag) included in production routes.web.php — protected by token but should be removed or env-gated entirely.')

add_heading(2, '8.3 TODO / FIXME Comments Related to Auth, Multi-Tenancy, AI')
simple_table(
    ['File', 'Line', 'Comment'],
    [
        ['app/Jobs/Agent/SubmitApplicationsJob.php', '168', 'TODO: Dispatch approval notification'],
        ['app/Jobs/Agent/SubmitApplicationsJob.php', '192', 'TODO: Dispatch success notification'],
        ['app/Jobs/Agent/FollowUpJob.php',           '128', 'TODO: Send notification to user'],
        ['app/Jobs/TrackEmployerBrandJob.php',        '386', 'TODO: Send email notification to company administrators'],
        ['app/Jobs/TrackEmployerBrandJob.php',        '399', 'TODO: Send email notification to company administrators'],
        ['app/Jobs/RefineLearningModelJob.php',       '403', 'TODO: Implement logic to get all active companies and dispatch jobs'],
        ['app/Jobs/GenerateBehavioralAssessmentJob.php', '284', 'TODO: Implement notification logic'],
        ['app/Jobs/GenerateBehavioralAssessmentJob.php', '315', 'TODO: Implement failure notification logic'],
        ['app/Services/Calendar/CalendarService.php', '103', 'TODO: Send cancellation notification'],
        ['app/Http/Controllers/ContactController.php', '24', 'TODO: Send email to support team'],
    ],
    [2.8, 0.6, 3.1]
)

add_heading(2, '8.4 Other Notable Gaps')
add_bullet('No GDPR cookie consent banner — GDPR service exists (GDPRService) but UI consent is not visible in marketing pages.')
add_bullet('No rate limiting on AI API routes beyond what Sanctum provides — high AI call volumes could exhaust Azure quota rapidly.')
add_bullet('Email templates: ContactController sends no email (TODO comment) — contact form is non-functional.')
add_bullet('Video interview transcription depends on Twilio — no fallback if Twilio is unavailable.')
add_bullet('Autonomous agent approval notifications are not implemented (TODO) — users may not know when agent needs approval.')
add_bullet('Meilisearch is the configured search driver but is optional — Scout falls back gracefully if not running.')

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — FILE STRUCTURE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '9. File Structure Summary')

simple_table(
    ['Directory', 'Purpose'],
    [
        ['app/',                   'All PHP application code'],
        ['app/Actions/',           'Single-purpose business actions (CreateResume, etc.)'],
        ['app/Console/',           'Artisan commands (orin:process-deadlines, ai:seed-disclaimers, etc.)'],
        ['app/Events/',            'Laravel events (JobApplied, ApplicationStatusChanged, etc.)'],
        ['app/Exceptions/',        'Custom exception classes'],
        ['app/Filament/',          'Filament admin panel: Resources, Pages, Widgets'],
        ['app/Helpers/',           'Global helper classes'],
        ['app/Http/Controllers/',  'Web + API controllers (grouped by domain: Admin/, API/, Auth/, Employer/, Marketplace/, Webhooks/)'],
        ['app/Http/Middleware/',   'Custom middleware (EnsureUserIsEmployer, ApiRateLimiting, SecurityHeaders, etc.)'],
        ['app/Jobs/',              '40+ queue jobs (Agent/, Marketplace/ sub-folders)'],
        ['app/Listeners/',         'Event listeners (HandleJobApplied, NotifyEmployer, etc.)'],
        ['app/Livewire/',          'Livewire components (SwipeJobBrowser, ProfileWizard, etc.)'],
        ['app/Mail/',              'Mailable classes (CandidateHiringMail, ApplicationConfirmationMail, etc.)'],
        ['app/Models/',            '200+ Eloquent models'],
        ['app/Notifications/',     'Laravel notification classes (27+ covering all hiring events)'],
        ['app/Observers/',         'Model observers (ApplicationObserver — cache busting)'],
        ['app/Policies/',          'Authorization policies (ResumePolicy, etc.)'],
        ['app/Providers/',         'Service providers (AppServiceProvider, FortifyServiceProvider, etc.)'],
        ['app/Services/',          'Business logic services (AI/, Agent/, Ats/, Calendar/, Interview/, JobBoard/, ResponsibleAI/, Search/, Subscription/ sub-folders)'],
        ['app/Traits/',            'Reusable PHP traits (InteractsWithAI, etc.)'],
        ['bootstrap/',             'Laravel bootstrap files (app.php, providers.php)'],
        ['config/',                '20+ config files (ai.php, payment.php, video-interview.php, etc.)'],
        ['database/migrations/',   '~140 migration files spanning Oct 2025 – May 2026'],
        ['database/seeders/',      'Database seeders (UserSeeder, SubscriptionPlanSeeder, etc.)'],
        ['database/factories/',    'Eloquent factories for testing'],
        ['public/',                'Web root (compiled assets, images, .user.ini for OPcache)'],
        ['resources/views/',       'Blade templates (layouts, components, employer/, resume/, career-coach/, etc.)'],
        ['resources/js/',          'Alpine.js + Vite entry points'],
        ['resources/css/',         'Tailwind CSS + custom styles'],
        ['routes/',                'Route files: web.php, api.php, auth.php, employer.php, resume.php, admin_analytics.php, console.php'],
        ['storage/',               'Logs, cache, uploaded files (Azure Blob in production)'],
        ['tests/',                 'PHPUnit / Pest tests (Feature/, Unit/)'],
        ['vendor/',                'Composer dependencies'],
    ],
    [2.0, 4.5]
)

add_heading(2, '9.1 Route Files')
simple_table(
    ['File', 'Purpose'],
    [
        ['routes/web.php',            'Main web routes: marketing, dashboard, jobs, career coach, payments, admin diag. Includes all sub-route files.'],
        ['routes/auth.php',           'Auth routes: login, register, password reset, 2FA, social OAuth (Google, LinkedIn, etc.)'],
        ['routes/employer.php',       'Employer portal: dashboard, ATS, talent pool, messaging, job wizard, referrals (all auth+employer middleware)'],
        ['routes/resume.php',         'Resume CRUD + AI features + public share view'],
        ['routes/api.php',            'All JSON API routes: Sanctum-auth routes, /api/v1/ third-party routes, skills, agent, negotiation, GDPR'],
        ['routes/admin_analytics.php','Admin analytics routes (auth:admin middleware)'],
        ['routes/console.php',        'Scheduled commands: Orin pipeline, agent jobs, skill analysis, talent pipelines, responsible AI bias checks'],
    ],
    [2.0, 4.5]
)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 10 — DEPENDENCIES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(1, '10. Dependencies')

add_heading(2, '10.1 PHP Dependencies (composer.json — Production)')
simple_table(
    ['Package', 'Version', 'Purpose', 'Flag'],
    [
        ['laravel/framework',            '^12.0',   'Core Laravel framework',                                              'OK'],
        ['filament/filament',            '^4.1',    'Admin panel + UI components',                                         'OK'],
        ['livewire/livewire',            '^3.8',    'Server-side reactive components',                                     'OK'],
        ['laravel/sanctum',              '^4.2',    'API token authentication',                                            'OK'],
        ['laravel/fortify',              '^1.31',   'Authentication actions (login, register, 2FA)',                       'OK'],
        ['laravel/socialite',            '^5.24',   'OAuth social login (Google, LinkedIn etc.)',                          'OK'],
        ['spatie/laravel-permission',    '^6.22',   'Role-based access control',                                          'OK'],
        ['spatie/laravel-medialibrary',  '^11.22',  'File uploads with conversions',                                      'OK'],
        ['spatie/laravel-data',          '^4.22',   'Typed DTOs for data transfer',                                       'OK'],
        ['spatie/laravel-query-builder', '^6.3',    'Filterable/sortable query builder for API',                          'OK'],
        ['spatie/laravel-settings',      '^3.8',    'DB-backed app settings',                                             'OK'],
        ['openai-php/laravel',           '^0.17.1', 'OpenAI SDK (used as legacy shim; primary calls use direct HTTP)',    'Review — primary AI uses direct HTTP to Azure; this package mostly unused'],
        ['laravel/horizon',              '^5.29',   'Redis queue monitoring dashboard',                                   'OK'],
        ['laravel/reverb',               '^1.10',   'WebSockets server for real-time features',                           'OK'],
        ['laravel/scout',                '^10.20',  'Full-text search abstraction (Meilisearch)',                         'OK'],
        ['meilisearch/meilisearch-php',  '^1.16',   'Meilisearch client',                                                 'OK'],
        ['predis/predis',                '^3.2',    'Redis PHP client',                                                   'OK'],
        ['razorpay/razorpay',            '^2.9',    'Razorpay payment gateway SDK',                                       'OK — verify live key rotation (see security alerts)'],
        ['stripe/stripe-php',            '^16.3',   'Stripe payment gateway SDK',                                         'OK'],
        ['barryvdh/laravel-dompdf',      '^3.1',    'PDF generation (resume export)',                                     'OK'],
        ['smalot/pdfparser',             '^2.12',   'PDF text extraction (resume upload parsing)',                        'OK'],
        ['intervention/image',           '^3.11',   'Image resizing/optimisation',                                        'OK'],
        ['matthewbdaly/laravel-azure-storage', '^2.0', 'Azure Blob Storage filesystem driver',                           'OK'],
        ['darkaonline/l5-swagger',       '^8.6',    'OpenAPI/Swagger documentation generator',                           'OK'],
        ['openspout/openspout',          '^4.28',   'Excel/CSV export (analytics, ATS export)',                          'OK'],
        ['sentry/sentry-laravel',        '^4.14',   'Error tracking and performance monitoring',                         'OK'],
        ['stancl/jobpipeline',           '^1.8',    'Chainable job pipeline',                                            'OK'],
        ['laravel/telescope',            '^5.20',   'Dev debugging dashboard (should be disabled in production)',        'Check TELESCOPE_ENABLED env is false in production'],
        ['laravel/tinker',               '^2.10.1', 'REPL for artisan (dev only)',                                       'Dev only — OK'],
    ],
    [2.2, 0.8, 2.4, 1.1]
)

add_heading(2, '10.2 PHP Dev Dependencies')
simple_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['fakerphp/faker',       '^1.23',   'Test data generation'],
        ['laravel/breeze',       '^2.3',    'Auth scaffolding (likely used during setup, may be vestigial)'],
        ['laravel/pail',         '^1.2.2',  'Real-time log streaming in terminal'],
        ['laravel/pint',         '^1.24',   'PHP code style fixer (PSR-12)'],
        ['laravel/sail',         '^1.41',   'Docker dev environment'],
        ['mockery/mockery',      '^1.6',    'Mock objects for tests'],
        ['nunomaduro/collision', '^8.6',    'Better CLI error display'],
        ['phpunit/phpunit',      '^11.5.3', 'Test framework'],
    ],
    [2.0, 1.0, 4.0]
)

add_heading(2, '10.3 JavaScript Dependencies (package.json)')
simple_table(
    ['Package', 'Version', 'Purpose'],
    [
        ['vite',                  '^7.0.7', 'Build tool'],
        ['laravel-vite-plugin',   '^2.0.0', 'Laravel Vite integration'],
        ['tailwindcss',           '^3.4.0', 'Utility CSS framework'],
        ['@tailwindcss/forms',    '^0.5.2', 'Tailwind form styles'],
        ['alpinejs',              '^3.4.2', 'Lightweight JS reactivity (used alongside Livewire)'],
        ['axios',                 '^1.11.0','HTTP client for browser AJAX calls'],
        ['autoprefixer',          '^10.4.2','CSS vendor prefix automation'],
        ['postcss',               '^8.4.31','CSS processing'],
        ['concurrently',          '^9.0.1', 'Run multiple dev processes simultaneously'],
    ],
    [2.0, 1.0, 4.0]
)

add_heading(2, '10.4 Flags for Orin Integration Compatibility')
add_bullet('laravel/breeze (dev): Likely vestigial — installed during initial scaffolding. Verify it does not conflict with Fortify auth routes.')
add_bullet('openai-php/laravel: Only used as a compatibility shim via OpenAIService. All real AI calls use direct HTTP to Azure. This creates a misleading dependency — consider removing the package if Orin exclusively uses Azure.')
add_bullet('laravel/telescope: Should confirm TELESCOPE_ENABLED=false in production — Telescope can add significant memory overhead.')
add_bullet('stancl/jobpipeline: Used for the Orin™ application pipeline. Ensure compatibility with Laravel 12 queue changes.')
add_bullet('laravel/reverb: WebSocket server — ensure Reverb is running in production if real-time features (notifications, live coaching) are active.')

# ─────────────────────────────────────────────────────────────────────────────
# SECURITY SUMMARY APPENDIX
# ─────────────────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(1, 'Appendix A — Security Issues Summary')

add_warning('P0 — CRITICAL: Azure OpenAI API key hardcoded in config/ai.php and AIService.php. Rotate immediately and use env-only.')
add_warning('P0 — CRITICAL: Live Razorpay API key (rzp_live_SYaE3xGta8COx9) shared in rzp-key CSV. Revoke and regenerate immediately via Razorpay dashboard.')
add_warning('P1 — HIGH: /auth-diag route is live in production routes.web.php. Although protected by AUTH_DIAG_TOKEN, it exposes schema introspection, user seeding, and migration triggering capabilities. Should be removed entirely from production.')
add_note('P2 — MEDIUM: openai-php/laravel package included but primary AI calls bypass it (direct HTTP). If OPENAI_API_KEY is set, the package auto-configures — ensure it is not accidentally used with a different key.')
add_note('P2 — MEDIUM: Phone numbers encrypted at rest (good), but verify Laravel encryption key (APP_KEY) rotation policy.')
add_note('P3 — LOW: Multiple TODO notifications mean some agent actions complete silently without user feedback.')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('— End of Report —')
run.font.color.rgb = MID_GREY
run.font.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
output_path = 'STUDAI_CAREER_TECHNICAL_AUDIT_REPORT.docx'
doc.save(output_path)
print(f'Report saved: {output_path}')
