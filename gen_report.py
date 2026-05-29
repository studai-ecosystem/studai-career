from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

BLUE  = RGBColor(0x1A, 0x73, 0xE8)
DARK  = RGBColor(0x1F, 0x29, 0x37)
MID   = RGBColor(0x4B, 0x55, 0x63)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def sf(run, size, bold=False, color=None, italic=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def h(text, size=10.5, color=BLUE, sb=5, sa=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    r = p.add_run(text)
    sf(r, size, bold=True, color=color)
    return p


def body(text, size=8.5, sa=2, color=MID):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    sf(r, size, color=color)
    return p


def fill_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def cell_pad(tbl, sb=1.5, sa=1.5):
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(sb)
                p.paragraph_format.space_after  = Pt(sa)


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A73E8')
    pBdr.append(bot)
    pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(1)
r1 = p.add_run('StudAI Career')
sf(r1, 20, bold=True, color=BLUE)
r2 = p.add_run('  \u00b7  Complete Product Report')
sf(r2, 11, color=MID)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(3)
r3 = p2.add_run(
    '"Your Career. On Autopilot."   \u2022   May 2026   \u2022   '
    'studai-app-prod.azurewebsites.net   \u2022   Laravel 12 / Azure'
)
sf(r3, 8, italic=True, color=MID)
divider()

# ══════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════
h('1. What Is StudAI Career?', size=10.5, sb=3, sa=1)
body(
    'StudAI Career is a fully deployed, AI-powered SaaS career platform for professionals at every stage. '
    'It unifies AI coaching, mock interviews, job search, resume analysis, salary negotiation, video '
    'interviews, mentorship, networking, gamification, payments, and employer-side hiring tools into one '
    'product. Built on Laravel 12 with Filament 4 admin, hosted on Azure App Service, and powered by OpenAI GPT-4o.'
)

# ══════════════════════════════════════════════════════════════
# SECTION 2 — ALL FEATURES  (3-column compact table)
# ══════════════════════════════════════════════════════════════
h('2. Complete Feature List', size=10.5, sb=4, sa=2)

categories = [
    ('AI & Intelligence', [
        'AI Mock Interviews (5-20 Qs, role-specific)',
        'AI Career Coach (multi-turn GPT-4o chat)',
        'AI Cover Letter Generator',
        'AI Skill Gap Analyzer',
        'AI Job Matching & scoring',
        'AI Application Assistant',
        'Salary Negotiation Trainer',
        'AI Performance Reports',
    ]),
    ('Resume & ATS', [
        'Resume Builder & Manager',
        'ATS Score Analyzer',
        'Resume Customizer per job',
        'Resume Export (PDF/Word)',
        'Resume-to-job fit scoring',
    ]),
    ('Job Marketplace', [
        'Job Search (Scout + Meilisearch)',
        'Advanced filters (title, location)',
        'Quick Apply (one-click)',
        'Swipe Job Browser (mobile)',
        'Application Tracker',
        'Public Apply (no login needed)',
        'Job Salary Benchmarks',
    ]),
    ('Video & Assessment', [
        'Video Interview Recorder',
        'Video transcription & AI analysis',
        'Candidate Assessments & Tests',
        'ATS Pipeline management',
        'Evaluation results & scoring',
    ]),
    ('Career Development', [
        'Career Profile & Goals',
        'Profile Wizard (onboarding)',
        'Mentorship Hub',
        'Networking & Connections',
        'Group Browser & Events (RSVP)',
        'Messaging Center (in-app)',
        'Activity Feed',
        'Gamification (badges, points)',
    ]),
    ('Company & Reviews', [
        'Company Profiles & Listings',
        'Company Reviews (submit/browse)',
        'Salary crowdsourcing',
        'Company Insights (AI)',
        'Interview experience sharing',
    ]),
    ('Scheduling & Calendar', [
        'Calendar & Availability mgmt',
        'Scheduling Links (Calendly-style)',
        'Event Browser & RSVP',
        'Interview Scheduling',
    ]),
    ('Payments & Plans', [
        'Subscription management',
        'Stripe & Razorpay gateways',
        'Webhook processing',
        'Payment history',
        'Offer Letter generation',
    ]),
    ('Analytics & Reporting', [
        'Application Funnel analytics',
        'Salary Benchmark charts',
        'Skills Forecast reports',
        'Heatmap visualizations',
        'Source Attribution tracking',
        'Time-to-Hire metrics',
        'Enhanced Analytics dashboard',
    ]),
    ('Interview Integrity', [
        'Fullscreen enforcement',
        'Tab-switch detection (3 strikes)',
        'Camera & Mic live monitor',
        'Anti-cheat auto-submit',
        'Per-question timer',
    ]),
    ('Platform & Security', [
        'Two-Factor Auth (TOTP)',
        'Password Security & audit',
        'GDPR Consent management',
        'Push Notifications (web)',
        'Background Check integration',
        'Audit logging',
        'Email templates (custom)',
        'Newsletter / marketing',
    ]),
    ('Admin & DevOps', [
        'Filament 4 Admin (full CRUD)',
        'Role-based access (Spatie)',
        'GitHub Actions CI/CD to Azure',
        'Laravel Horizon queue workers',
        'Auto database migrations',
        'Health check endpoints',
        'Marketing pages (pricing/blog)',
        'PWA / Offline support',
    ]),
]

col1 = categories[0:4]
col2 = categories[4:8]
col3 = categories[8:12]


def write_cat_block(cell, cat_list):
    cell.text = ''
    for cat_name, items in cat_list:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(0.5)
        r = p.add_run(cat_name)
        sf(r, 7.5, bold=True, color=BLUE)
        for item in items:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(0.3)
            p2.paragraph_format.left_indent  = Cm(0.15)
            r2 = p2.add_run('\u2022  ' + item)
            sf(r2, 6.8, color=MID)


feat_tbl = doc.add_table(rows=1, cols=3)
feat_tbl.style = 'Table Grid'
fcells = feat_tbl.rows[0].cells

write_cat_block(fcells[0], col1)
write_cat_block(fcells[1], col2)
write_cat_block(fcells[2], col3)

for fc in fcells:
    fill_cell(fc, 'EEF4FF')

feat_tbl.columns[0].width = Cm(5.6)
feat_tbl.columns[1].width = Cm(5.6)
feat_tbl.columns[2].width = Cm(5.6)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════
# SECTION 3 — CONFIRMED WORKING
# ══════════════════════════════════════════════════════════════
h('3. Confirmed Working Features  (Production \u2014 May 2026)', size=10.5, sb=3, sa=2)

working = [
    ('\u2705', 'AI Mock Interviews',          'Full flow: session create \u2192 AI Qs \u2192 answer \u2192 submit \u2192 report. Production confirmed.'),
    ('\u2705', 'Camera & Microphone',          'nginx Permissions-Policy fixed. Camera/mic auto-activates on interview start.'),
    ('\u2705', 'Interview Integrity Mode',     'Fullscreen lock, tab-switch detection (3 strikes = auto-submit), per-question timer.'),
    ('\u2705', 'AI Career Coach',              'Multi-turn GPT-4o chat, scroll fix, voice error handling all confirmed.'),
    ('\u2705', 'Job Search & Marketplace',     'Search, filters, listings, Quick Apply, Scout + Meilisearch sync operational.'),
    ('\u2705', 'User Auth & Roles',            'Login, register, 2FA (TOTP + recovery codes), role-based access stable.'),
    ('\u2705', 'Resume Management',            'Upload, view, manage resumes. ATS analysis pipeline active via queue jobs.'),
    ('\u2705', 'Admin Panel (Filament 4)',     'Full CRUD, user management, interview session oversight, role-protected.'),
    ('\u2705', 'Azure CI/CD Pipeline',         'GitHub Actions \u2192 Azure App Service. Startup script stable. OOM issue resolved.'),
    ('\u2705', 'Auto DB Migrations',           'Idempotent migrations on every deploy. Schema guards prevent column crashes.'),
    ('\u2705', 'Queue Workers (Horizon)',      'Background AI jobs (resume analysis, report gen) via Horizon + Redis.'),
    ('\u2705', 'Payments (Stripe/Razorpay)',  'Payment gateways integrated, webhooks registered, subscription flows built.'),
    ('\u2705', 'Marketing & Public Pages',     'Welcome, pricing, features, about, blog, contact, terms, privacy all live.'),
    ('\u2705', 'Gamification',                 'Badge system, points, and activity feed components in place.'),
    ('\u2705', 'GDPR & Push Notifications',   'Consent capture, audit, data flows + web push notification service active.'),
]

wtbl = doc.add_table(rows=1, cols=3)
wtbl.style = 'Table Grid'
hdr = wtbl.rows[0].cells
for i, lbl in enumerate(['\u00a0', 'Feature', 'Production Status']):
    hdr[i].text = ''
    r = hdr[i].paragraphs[0].add_run(lbl)
    sf(r, 8, bold=True, color=WHITE)
    fill_cell(hdr[i], '1A73E8')

for icon, feat, notes in working:
    row = wtbl.add_row().cells
    row[0].text = ''
    r0 = row[0].paragraphs[0].add_run(icon)
    sf(r0, 8)
    row[1].text = ''
    r1 = row[1].paragraphs[0].add_run(feat)
    sf(r1, 7.5, bold=True, color=DARK)
    row[2].text = ''
    r2 = row[2].paragraphs[0].add_run(notes)
    sf(r2, 7, color=MID)

wtbl.columns[0].width = Cm(0.5)
wtbl.columns[1].width = Cm(4.0)
wtbl.columns[2].width = Cm(12.3)
cell_pad(wtbl, sb=1.2, sa=1.2)

doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ══════════════════════════════════════════════════════════════
# SECTION 4 — TECH STACK
# ══════════════════════════════════════════════════════════════
h('4. Technology Stack', size=10.5, sb=3, sa=2)

stack_tbl = doc.add_table(rows=2, cols=4)
stack_tbl.style = 'Table Grid'

stack_data = [
    [
        ('Backend', 'Laravel 12 \u00b7 PHP 8.3 \u00b7 MySQL (Azure)'),
        ('Frontend', 'Blade \u00b7 Livewire \u00b7 Alpine.js \u00b7 Tailwind \u00b7 Vite'),
        ('AI Engine', 'OpenAI GPT-4o \u00b7 openai-php/laravel \u00b7 Centralised AIService'),
        ('Admin', 'Filament 4 \u00b7 Spatie Permissions (RBAC)'),
    ],
    [
        ('Search', 'Laravel Scout \u00b7 Meilisearch'),
        ('Auth', 'Fortify \u00b7 Sanctum \u00b7 TOTP 2FA \u00b7 API Tokens'),
        ('Infra', 'Azure App Service (Linux) \u00b7 nginx \u00b7 GitHub Actions'),
        ('Queues', 'Laravel Horizon \u00b7 Redis \u00b7 Stripe/Razorpay Webhooks'),
    ],
]

for ri, row_data in enumerate(stack_data):
    for ci, (label, value) in enumerate(row_data):
        cell = stack_tbl.rows[ri].cells[ci]
        cell.text = ''
        p1 = cell.paragraphs[0]
        rl = p1.add_run(label + '\n')
        sf(rl, 7.5, bold=True, color=BLUE)
        rv = p1.add_run(value)
        sf(rv, 7, color=MID)
        fill_cell(cell, 'EEF4FF')
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after  = Pt(2)

for col in stack_tbl.columns:
    col.width = Cm(4.2)

divider()

pf = doc.add_paragraph()
pf.paragraph_format.space_before = Pt(2)
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run(
    'StudAI Career  \u00b7  Confidential Product Report  \u00b7  May 2026  '
    '\u00b7  studai-app-prod.azurewebsites.net'
)
sf(rf, 7, italic=True, color=MID)

out = r'C:\Users\user\Downloads\studai-career\StudAI_Career_Product_Report.docx'
doc.save(out)
print('Done:', out)
