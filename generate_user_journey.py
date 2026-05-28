"""
StudAI Hire — Complete User Journey Document (Formal, Updated with Screening)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLUE      = RGBColor(0x1A, 0x73, 0xE8)
DARK_BLUE = RGBColor(0x0D, 0x47, 0xA1)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG  = RGBColor(0xF1, 0xF5, 0xFE)
MID_BG    = RGBColor(0xD2, 0xE3, 0xFC)
GREY      = RGBColor(0x55, 0x55, 0x55)
GREEN     = RGBColor(0x18, 0x8A, 0x50)
RED_C     = RGBColor(0xB0, 0x22, 0x22)
AMBER     = RGBColor(0xD8, 0x6B, 0x00)

def cell_bg(cell, rgb):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), str(rgb).upper())
    tcPr.append(shd)

def para(doc, text, bold=False, size=10.5, colour=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if colour: r.font.color.rgb = colour
    return p

def divider(doc):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Inches(6.6)
    cell_bg(t.rows[0].cells[0], DARK_BLUE)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(1)
    t.rows[0].cells[0].paragraphs[0].paragraph_format.space_after  = Pt(1)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def section_title(doc, number, title):
    divider(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f'{number}  ')
    r1.bold = True; r1.font.size = Pt(13); r1.font.color.rgb = BLUE
    r2 = p.add_run(title.upper())
    r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = DARK_BLUE

def sub_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = BLUE

def flow_table(doc, steps):
    tbl = doc.add_table(rows=len(steps)+1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(1.4)
    tbl.columns[1].width = Inches(3.3)
    tbl.columns[2].width = Inches(1.9)
    for ci, h in enumerate(['Phase', 'User Action / System Step', 'Outcome / Screen']):
        cell_bg(tbl.rows[0].cells[ci], DARK_BLUE)
        p = tbl.rows[0].cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
    for i, (phase, action, outcome) in enumerate(steps, 1):
        row = tbl.rows[i]
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        cell_bg(row.cells[0], MID_BG)
        pp = row.cells[0].paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.space_before = Pt(3); pp.paragraph_format.space_after = Pt(3)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        rp = pp.add_run(phase); rp.bold = True; rp.font.size = Pt(8.5); rp.font.color.rgb = DARK_BLUE
        cell_bg(row.cells[1], bg)
        pa = row.cells[1].paragraphs[0]
        pa.paragraph_format.space_before = Pt(3); pa.paragraph_format.space_after = Pt(3)
        ra = pa.add_run(action); ra.font.size = Pt(9.5)
        cell_bg(row.cells[2], bg)
        po = row.cells[2].paragraphs[0]
        po.paragraph_format.space_before = Pt(3); po.paragraph_format.space_after = Pt(3)
        ro = po.add_run(outcome); ro.font.size = Pt(9); ro.italic = True; ro.font.color.rgb = GREY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def decision_table(doc, question, yes_path, no_path, yes_label='YES  ✓', no_label='NO   ✗'):
    dt = doc.add_table(rows=1, cols=1)
    dt.style = 'Table Grid'; dt.alignment = WD_TABLE_ALIGNMENT.LEFT
    dt.columns[0].width = Inches(6.6)
    cell_bg(dt.rows[0].cells[0], BLUE)
    dq = dt.rows[0].cells[0].paragraphs[0]
    dq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dq.paragraph_format.space_before = Pt(3); dq.paragraph_format.space_after = Pt(3)
    rq = dq.add_run(f'◆  DECISION:  {question}')
    rq.bold = True; rq.font.size = Pt(9.5); rq.font.color.rgb = WHITE
    bt = doc.add_table(rows=1, cols=2)
    bt.style = 'Table Grid'; bt.alignment = WD_TABLE_ALIGNMENT.LEFT
    bt.columns[0].width = Inches(3.3); bt.columns[1].width = Inches(3.3)
    for ci, (label, steps, hdr_col, bg) in enumerate([
        (yes_label, yes_path, GREEN,  LIGHT_BG),
        (no_label,  no_path,  RED_C,  RGBColor(0xFF, 0xF0, 0xF0)),
    ]):
        cell = bt.rows[0].cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        cell_bg(cell, bg)
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.space_before = Pt(3); hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run(label); hr.bold = True; hr.font.size = Pt(9.5); hr.font.color.rgb = hdr_col
        for step in steps:
            sp = cell.add_paragraph(f'  • {step}')
            sp.paragraph_format.space_before = Pt(1); sp.paragraph_format.space_after = Pt(1)
            sp.runs[0].font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def pipeline_table(doc, stages):
    """Horizontal pipeline: Stage → Status → Triggered By → Candidate Notified?"""
    tbl = doc.add_table(rows=len(stages)+1, cols=4)
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(1.5)
    tbl.columns[2].width = Inches(2.1)
    tbl.columns[3].width = Inches(1.5)
    for ci, h in enumerate(['Stage', 'ATS Status', 'Triggered By', 'Candidate Notified']):
        cell_bg(tbl.rows[0].cells[ci], DARK_BLUE)
        p = tbl.rows[0].cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WHITE
    for i, (stage, status, trigger, notified) in enumerate(stages, 1):
        row = tbl.rows[i]
        bg = LIGHT_BG if i % 2 == 0 else WHITE
        for ci in range(4): cell_bg(row.cells[ci], bg)
        row.cells[0].paragraphs[0].add_run(stage).font.size  = Pt(9)
        row.cells[1].paragraphs[0].add_run(status).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(trigger).font.size = Pt(9)
        n_run = row.cells[3].paragraphs[0].add_run(notified)
        n_run.font.size = Pt(9)
        n_run.font.color.rgb = GREEN if 'Yes' in notified else RED_C
        for ci in range(4):
            row.cells[ci].paragraphs[0].paragraph_format.space_before = Pt(3)
            row.cells[ci].paragraphs[0].paragraph_format.space_after  = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ═══════════════════════════════════════════════════════════════
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.85); sec.bottom_margin = Inches(0.85)
sec.left_margin = Inches(1.0);  sec.right_margin  = Inches(1.0)

# ── COVER PAGE ────────────────────────────────────────────────
bar = doc.add_table(rows=1, cols=1)
bar.style = 'Table Grid'; bar.columns[0].width = Inches(6.6); bar.alignment = WD_TABLE_ALIGNMENT.LEFT
cell_bg(bar.rows[0].cells[0], DARK_BLUE)
bp = bar.rows[0].cells[0].paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(18); bp.paragraph_format.space_after = Pt(18)
br = bp.add_run('STUDAI HIRE'); br.bold = True; br.font.size = Pt(28); br.font.color.rgb = WHITE
para(doc, '')
p_tag = doc.add_paragraph()
p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_tag.paragraph_format.space_before = Pt(4); p_tag.paragraph_format.space_after = Pt(2)
rt = p_tag.add_run('Your Career. On Autopilot.')
rt.font.size = Pt(12); rt.italic = True; rt.font.color.rgb = GREY
para(doc, ''); para(doc, '')
p_main = doc.add_paragraph()
p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_main.paragraph_format.space_before = Pt(8); p_main.paragraph_format.space_after = Pt(4)
rm = p_main.add_run('User Journey & Product Flow')
rm.bold = True; rm.font.size = Pt(20); rm.font.color.rgb = DARK_BLUE
para(doc, '')
meta = [
    ('Document Type', 'Product User Journey — Complete Edition'),
    ('Product',       'StudAI Hire'),
    ('Version',       'Phase 4 — Production'),
    ('Date',          'May 28, 2026'),
    ('Audience',      'Product  ·  Design  ·  QA  ·  Stakeholders'),
]
mt = doc.add_table(rows=len(meta), cols=2)
mt.style = 'Table Grid'; mt.alignment = WD_TABLE_ALIGNMENT.CENTER
mt.columns[0].width = Inches(2.0); mt.columns[1].width = Inches(3.5)
for i, (k, v) in enumerate(meta):
    kc = mt.rows[i].cells[0]; vc = mt.rows[i].cells[1]
    cell_bg(kc, DARK_BLUE)
    kp = kc.paragraphs[0]; kp.paragraph_format.space_before = Pt(3); kp.paragraph_format.space_after = Pt(3)
    kr = kp.add_run(k); kr.bold = True; kr.font.size = Pt(9.5); kr.font.color.rgb = WHITE
    vp = vc.paragraphs[0]; vp.paragraph_format.space_before = Pt(3); vp.paragraph_format.space_after = Pt(3)
    vr = vp.add_run(v); vr.font.size = Pt(9.5)
doc.add_page_break()

# ── 1. OVERVIEW ───────────────────────────────────────────────
section_title(doc, '1.', 'Platform Overview')
para(doc,
    'StudAI Hire is an AI-powered recruitment and career development SaaS platform. '
    'It serves three user roles, each with a structured journey across discovery, '
    'profile-building, screening, hiring, and post-hire workflows.',
    size=10, before=4, after=6)
roles_data = [
    ('Job Seeker',
     'Registers, builds a profile, searches jobs, applies, completes AI-driven screening tests, '
     'attends interviews, and tracks their full hiring pipeline with AI tools support.'),
    ('Employer',
     'Onboards a company, posts jobs via AI Wizard, screens candidates through Orin™ evaluation, '
     'manages a full ATS pipeline, schedules interviews, requests background checks, and analyses hiring metrics.'),
    ('Admin',
     'Manages platform operations via the Filament admin panel — users, subscriptions, '
     'job moderation, AI credits, and system health.'),
]
rt2 = doc.add_table(rows=len(roles_data)+1, cols=2)
rt2.style = 'Table Grid'; rt2.alignment = WD_TABLE_ALIGNMENT.LEFT
rt2.columns[0].width = Inches(1.4); rt2.columns[1].width = Inches(5.2)
for ci, h in enumerate(['Role', 'Responsibilities']):
    cell_bg(rt2.rows[0].cells[ci], DARK_BLUE)
    p = rt2.rows[0].cells[ci].paragraphs[0]
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = WHITE
for i, (role, desc) in enumerate(roles_data, 1):
    bg = LIGHT_BG if i % 2 == 0 else WHITE
    for ci in range(2): cell_bg(rt2.rows[i].cells[ci], bg)
    rp = rt2.rows[i].cells[0].paragraphs[0]
    rp.paragraph_format.space_before = Pt(4); rp.paragraph_format.space_after = Pt(4)
    rt2.rows[i].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    rr = rp.add_run(role); rr.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = DARK_BLUE
    dp = rt2.rows[i].cells[1].paragraphs[0]
    dp.paragraph_format.space_before = Pt(4); dp.paragraph_format.space_after = Pt(4)
    dr = dp.add_run(desc); dr.font.size = Pt(9)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# §2  JOB SEEKER JOURNEY
# ═══════════════════════════════════════════════════════════════
section_title(doc, '2.', 'Job Seeker Journey')

sub_title(doc, '2.1  Discovery & Registration')
flow_table(doc, [
    ('Discovery',    'Lands on studai.one — browses homepage, features, pricing, and how-it-works.', '/  /features  /pricing'),
    ('Registration', 'Clicks "Get Started Free" — enters name, email, and password.',               '/register'),
    ('Verification', 'Receives and clicks the email verification link to activate the account.',     'Email → /verify-email'),
    ('Social Login', 'Alternatively authenticates via Google or LinkedIn OAuth 2.0.',                'SocialAuthController'),
    ('Dashboard',    'Authenticated and redirected to the personalised job seeker dashboard.',       '/dashboard'),
])

sub_title(doc, '2.2  Profile & Subscription Setup')
flow_table(doc, [
    ('Profile Hub',  'Navigates to Career Profile to unlock AI-powered job matching.',               '/profile/career'),
    ('Resume Upload','Uploads existing resume (PDF/DOCX) — AI auto-fills all sections.',             'AI parse + auto-fill'),
    ('Manual Build', 'Alternatively completes the step-by-step Profile Builder Wizard.',            '/profile/career/builder'),
    ('AI Suggestions','AI analyses profile and recommends targeted improvements.',                   'CareerProfileController'),
    ('Plan Selection','Compares Free / Pro / Elite plan tiers.',                                     '/pricing'),
    ('Payment',      'Completes payment via Razorpay (India) or Stripe (International).',           'Secure checkout'),
    ('Activation',   'Subscription activated; AI credits granted; all premium features unlocked.',   '/dashboard'),
])

sub_title(doc, '2.3  Job Search & Application')
flow_table(doc, [
    ('Search',       'Searches jobs by role, location, salary, type, and remote preference.',        '/jobs/search'),
    ('Job Detail',   'Reviews full job description, requirements, salary range, and company info.',  '/jobs/{id}'),
    ('Save Job',     'Bookmarks a listing for later review.',                                         '/jobs/saved'),
    ('Cover Letter', 'Optionally requests an AI-generated, tailored cover letter.',                  'AI — 1 credit'),
    ('Apply',        'Submits application with selected resume and cover letter.',                    'Application created'),
    ('Confirmation', 'Status set to "Applied"; employer notified automatically via queue.',           '/applications'),
])

sub_title(doc, '2.4  Orin™ AI Screening (Public Apply Link)')
para(doc,
    'When an employer shares a unique /apply/{token} link, candidates are routed through '
    'a structured, AI-proctored evaluation before their application is accepted.',
    size=9.5, italic=True, colour=GREY, before=0, after=4)
flow_table(doc, [
    ('Landing',      'Candidate opens the employer-shared link; sees job title, company, and round overview.', '/apply/{token}'),
    ('Submit Info',  'Fills basic details and uploads resume.',                                       'POST /apply/{token}/submit'),
    ('Eval. Portal', 'Redirected to the evaluation interface.',                                       '/apply/{token}/evaluation'),
    ('Round 1',      'Completes Company Information Test — multiple-choice questions.',               'OrinEvaluationService'),
    ('Round 2',      'Completes Aptitude / Coding / Case Study round (if configured).',              'Timed, proctored'),
    ('Anti-Cheat',   'System records tab-switch, copy-paste, and focus-loss events silently.',       'Anti-cheat API'),
    ('Scoring',      'AI scores each answer in real time; calculates an overall evaluation score.',  'OrinEvaluationService'),
    ('Results',      'Candidate sees their score, AI feedback, and next-step instructions.',         '/apply/{token}/results'),
    ('ATS Entry',    'Application with evaluation score enters the employer ATS automatically.',      'ATS pipeline'),
])

sub_title(doc, '2.5  Employer-Assigned Hiring Tests (Post-Shortlist)')
para(doc,
    'Once shortlisted inside the ATS, candidates receive a secure token link by email to '
    'complete stage-specific tests set by the employer.',
    size=9.5, italic=True, colour=GREY, before=0, after=4)
flow_table(doc, [
    ('Email Link',   'Candidate receives a stage-specific test link via email.',                      'HiringTestController email'),
    ('Test Page',    'Opens the timed test — instructions, time limit, and questions displayed.',    '/hiring-test/{token}/{stage}'),
    ('Stage: Info',  'Answers company-culture / domain-knowledge MCQ questions.',                    'Stage: company_info_test'),
    ('Stage: Apt.',  'Answers aptitude / logical reasoning questions.',                               'Stage: aptitude'),
    ('Submit',       'Submits answers before the time limit expires.',                               'POST /hiring-test/{token}/{stage}/submit'),
    ('Scoring',      'System auto-grades the test; pass/fail threshold applied.',                    'Pass score configured by employer'),
    ('Result',       'Employer notified; test result updates the applicant rank in ATS.',            'Rank score updated'),
])

doc.add_page_break()

sub_title(doc, '2.6  Video Interview (Self-Practice & Invited)')
flow_table(doc, [
    ('Practice Hub', 'Accesses the Video Interview studio for self-practice sessions.',               '/video-interview'),
    ('Create Mock',  'Sets topic, difficulty, and number of questions for a practice session.',      '/video-interview/create'),
    ('Record',       'Records video responses via browser webcam — Livewire-powered.',               '/video-interview/record/{session}'),
    ('AI Analysis',  'AI analyses tone, pacing, vocabulary, and confidence from the recording.',     'VideoAnalysisService'),
    ('Results',      'Receives a detailed breakdown with scores and improvement suggestions.',        '/video-interview/results/{session}'),
    ('Employer Inv.','Receives an employer-sent video interview invitation link via email.',          '/video-interview/invitation/{token}'),
    ('Accept/Decline','Accepts or declines the employer invitation.',                                'POST /invitation/{id}/accept'),
    ('Complete',     'Records and uploads the employer interview response.',                          'VideoInterviewController::uploadVideo()'),
])

sub_title(doc, '2.7  Application Status Tracking & Outcome')
flow_table(doc, [
    ('Dashboard',    'Tracks all applications in real time with colour-coded status badges.',        '/applications'),
    ('Notification', 'Receives in-app and email notification on every pipeline stage change.',       'CandidateHiringMail'),
    ('Shortlisted',  'Notified of shortlisting; may be prompted to complete a hiring test.',         'CandidateShortlistedNotification'),
    ('Interviewed',  'Interview scheduled — date, time, and meeting link shared by employer.',       'Calendar invite'),
    ('Offer',        'Receives formal offer letter when hired; status updated to Hired.',             'CandidateHiredNotification'),
    ('Rejected',     'Receives a respectful rejection; Career Coach suggests next steps.',           'CandidateRejectedNotification'),
])

sub_title(doc, '2.8  AI Career Tools')
flow_table(doc, [
    ('Resume Builder', 'Creates ATS-optimised resumes with AI summary, skill extraction, and export.', '/resume'),
    ('ATS Checker',   'Runs an AI ATS score check against a target job description.',                '/resume/{id}/ats-check'),
    ('Cover Letter',  'AI generates personalised cover letters; exported as PDF or DOCX.',           '/resume/{id}/cover-letter'),
    ('Mock Interview','Practises with AI-generated interview questions; receives scored feedback.',   '/interview/session/{id}'),
    ('Skill Analyzer','Identifies skill gaps; follows curated learning paths and daily modules.',     '/skills/dashboard'),
    ('Career Coach',  'Conducts AI chat coaching sessions; sets and tracks career goals weekly.',    '/career-coach'),
    ('Auto-Apply Agent','Activates autonomous AI agent to scan, score, and apply to jobs 24/7.',    '/agent/dashboard'),
    ('Negotiation',   'Receives AI salary negotiation script with range data and tactics.',          '/interview/salary-negotiation'),
    ('Company Reviews','Reads verified peer reviews, salary reports, and interview experiences.',    '/companies/{slug}/reviews'),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# §3  EMPLOYER JOURNEY
# ═══════════════════════════════════════════════════════════════
section_title(doc, '3.', 'Employer Journey')

sub_title(doc, '3.1  Registration & Company Onboarding')
flow_table(doc, [
    ('Registration', 'Signs up at /register with account type set to Employer.',                    '/register'),
    ('Verification', 'Confirms email address via the verification link.',                            'Email → /verify-email'),
    ('Onboarding',   'Completes Company Wizard — name, logo, size, industry, culture description.', '/employer/onboarding'),
    ('AI Assist',    'AI auto-generates company description from name and industry context.',         'AI suggestion'),
    ('Dashboard',    'Employer Dashboard activated — all hiring tools accessible.',                  '/employer'),
])

sub_title(doc, '3.2  Job Posting with AI Wizard')
flow_table(doc, [
    ('Wizard',       'Opens the AI-powered Job Posting Wizard.',                                     '/employer/jobs/wizard/start'),
    ('Details',      'Enters job title, department, location, employment type.',                     'Step 1'),
    ('AI Job Desc.', 'AI generates a complete, professional job description from the input.',         'Step 2 — AI draft'),
    ('Requirements', 'Defines required skills, experience, salary range, and team.',                 'Step 3'),
    ('Pipeline',     'Configures the hiring pipeline stages and evaluation rounds.',                  'Step 4'),
    ('Publish',      'Reviews, previews, and publishes the job listing.',                             'Job live on platform'),
    ('Orin™ Link',  'Receives a unique /apply/{token} link to share with external candidates.',      'Shareable screening link'),
    ('Test Setup',   'Optionally creates stage-specific MCQ tests for shortlisted candidates.',      '/employer/jobs/{id}/tests/{stage}/create'),
])

sub_title(doc, '3.3  Screening & ATS Pipeline')
para(doc,
    'The full candidate pipeline is managed in the Applicant Tracking System (ATS). '
    'Every stage change triggers an automatic, AI-composed notification to the candidate.',
    size=9.5, italic=True, colour=GREY, before=0, after=4)

pipeline_table(doc, [
    ('Application Received', 'applied',      'Candidate submits via platform or /apply/{token}', 'Yes — confirmation'),
    ('Under Review',         'reviewing',    'Employer opens and views the application',          'No'),
    ('AI Pre-Screening',     'screened',     'Orin™ evaluation score assessed automatically',     'No'),
    ('Shortlisted',          'shortlisted',  'Employer moves candidate to shortlist',             'Yes — shortlist email'),
    ('Company Info Test',    'testing',      'Token link emailed to candidate for MCQ stage',     'Yes — test link'),
    ('Aptitude Test',        'testing',      'Second test stage dispatched after first passes',   'Yes — test link'),
    ('Interview Scheduled',  'interview',    'Employer schedules interview in system',            'Yes — calendar invite'),
    ('Offer Extended',       'offer',        'Employer issues offer letter',                      'Yes — offer email'),
    ('Hired',                'hired',        'Candidate accepts offer',                           'Yes — hired email'),
    ('Rejected',             'rejected',     'Employer rejects at any stage',                     'Yes — rejection email'),
])

sub_title(doc, '3.4  Interview Management')
flow_table(doc, [
    ('Interview Hub', 'Views all scheduled and pending interviews across all job postings.',         '/employer/interviews'),
    ('Schedule',      'Selects shortlisted candidate and opens scheduling form.',                    '/employer/interviews/schedule/{applicationId}'),
    ('Configure',     'Sets type (phone/video/onsite/panel), date, time, duration, meeting link.',  'InterviewManagementController::store()'),
    ('Notify',        'Candidate automatically notified with full interview details.',               'Queued notification'),
    ('Panel Scores',  'Each panel interviewer submits structured scores post-interview.',            'POST /employer/interviews/{id}/scores'),
    ('Evaluation',    'Hiring manager reviews panel scores and submits a final evaluation.',         '/employer/interviews/{id}/evaluate'),
    ('Decision',      'Makes final hire / reject decision from the decision form.',                  '/employer/interviews/{id}/decide'),
    ('Complete',      'Interview marked complete; ATS status updated accordingly.',                  'PATCH /employer/interviews/{id}/complete'),
])

sub_title(doc, '3.5  Hiring Tests — Employer Configuration')
flow_table(doc, [
    ('Test Creator',  'Opens the test creation form for a specific job and hiring stage.',           '/employer/jobs/{jobId}/tests/{stage}/create'),
    ('Configure',     'Defines title, instructions, time limit, pass score, and questions.',         'HiringTestManagerController::store()'),
    ('Question Types','Adds MCQ questions with options and correct-answer index.',                   'Min 1 question required'),
    ('Activate',      'Test saved and activated — candidates in that stage receive the link.',       'is_active = true'),
    ('Results View',  'Views all candidate test attempts, scores, and pass/fail outcomes.',          '/employer/jobs/{jobId}/tests/{stage}/results'),
])

sub_title(doc, '3.6  Background Checks')
flow_table(doc, [
    ('Initiate',      'After a conditional offer, employer initiates a background check.',           '/employer/background-checks/create'),
    ('Consent',       'Consent request emailed to candidate; candidate must digitally consent.',     'POST /background-checks/{id}/send-consent'),
    ('Verification',  'Background check service verifies identity, employment, and education.',      'BackgroundCheckService'),
    ('Report',        'Employer views the completed report.',                                        '/employer/background-checks/{id}'),
    ('Download',      'Full background check report downloaded as PDF for records.',                 '/employer/background-checks/{id}/download'),
    ('Hire/Withdraw', 'Final hiring decision confirmed or offer withdrawn based on results.',         'ATS updated'),
])

sub_title(doc, '3.7  Post-Hire & Talent Strategy')
flow_table(doc, [
    ('Offer Letter',  'Generates and emails a formal offer letter to the hired candidate.',         'OfferLetterController'),
    ('Talent Pool',   'Saves strong candidates for future openings; sends proactive outreach.',      '/employer/talent-pool'),
    ('Messaging',     'Sends personalised messages or bulk outreach via the messaging centre.',      '/employer/messages'),
    ('Referrals',     'Manages an employee referral programme with leaderboard and rewards.',        '/employer/referrals'),
    ('Analytics',     'Reviews hiring KPIs — time-to-hire, source attribution, diversity.',         '/employer/analytics'),
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# §4  ADMIN JOURNEY
# ═══════════════════════════════════════════════════════════════
section_title(doc, '4.', 'Admin Journey')
para(doc,
    'Platform administrators access a dedicated Filament 4.x admin panel. '
    'All operations are authenticated, role-gated, and audit-logged.',
    size=10, before=4, after=6)
flow_table(doc, [
    ('Login',        'Authenticates at the dedicated admin portal.',                                '/studai/login'),
    ('User Mgmt.',   'Views, bans, and impersonates users; assigns roles and permissions.',         'Filament panel'),
    ('Subscriptions','Creates and edits subscription plans, pricing tiers, and feature flags.',      'Plan management'),
    ('Job Moderation','Approves, features, or removes job listings from the platform.',             'Job listings resource'),
    ('AI Credits',   'Monitors AI credit consumption per user; adjusts limits as required.',        'ai_credit_logs'),
    ('Applications', 'Monitors all platform applications; resolves disputes and escalations.',       '/admin/application-monitor'),
    ('System Health','Checks queue status via Horizon, error logs, and the /up health endpoint.',   'Laravel Horizon'),
])

# ═══════════════════════════════════════════════════════════════
# §5  KEY SYSTEM FLOWS
# ═══════════════════════════════════════════════════════════════
section_title(doc, '5.', 'Key System Flows')

sub_title(doc, '5.1  Full Screening & Hiring Pipeline (End-to-End)')
para(doc,
    'The diagram below shows the complete screening funnel from application to hire, '
    'covering both the Orin™ public evaluation path and the internal ATS hiring test path.',
    size=9.5, italic=True, colour=GREY, before=0, after=4)
flow_table(doc, [
    ('Apply',        'Candidate submits application (platform or /apply/{token} link).',            'Application created'),
    ('AI Pre-Screen','Orin™ scores evaluation answers; result attached to application record.',      'OrinEvaluationService'),
    ('ATS Review',   'Employer reviews resume, cover letter, and Orin™ score in ATS.',             '/employer/ats'),
    ('Shortlist',    'Employer shortlists — candidate notified; stage-specific test link emailed.', 'Email dispatch'),
    ('Test Stage 1', 'Candidate completes Company Info Test (MCQ) within time limit.',              '/hiring-test/{token}/company_info_test'),
    ('Test Stage 2', 'Candidate completes Aptitude Test (if configured) within time limit.',        '/hiring-test/{token}/aptitude'),
    ('Score Check',  'System compares score vs. pass threshold; result recorded in ATS.',           'Auto-grading'),
    ('Interview',    'Employer schedules interview (phone / video / panel / onsite).',              '/employer/interviews/schedule'),
    ('Panel Eval.',  'Each interviewer submits scores; hiring manager submits final evaluation.',    'InterviewManagementController'),
    ('Decision',     'Final hire / reject decision submitted; ATS status updated.',                  '/employer/interviews/{id}/decide'),
    ('BG Check',     'Conditional offer issued; background check initiated and consent sent.',      'BackgroundCheckController'),
    ('Hire',         'BG check passed; formal offer letter issued; candidate status = Hired.',      'OfferLetterController'),
])

sub_title(doc, '5.2  Notification Flow')
flow_table(doc, [
    ('Trigger',   'An ATS stage change or hiring event occurs.',                                    'Eloquent event'),
    ('Queue',     'Notification job dispatched asynchronously to the notifications queue.',          'ShouldQueue'),
    ('Email',     'AI-composed personalised email delivered to candidate and/or HR.',               'CandidateHiringMail / HRHiringMail'),
    ('In-App',    'Bell icon updated; user reads notification at /notifications.',                  '/notifications'),
])

sub_title(doc, '5.3  AI Credit Flow')
flow_table(doc, [
    ('Request',   'User triggers an AI feature (resume summary, cover letter, coaching, etc.).',   'Any AI action'),
    ('Validate',  'System checks the user has a sufficient credit balance.',                         'Credit check'),
    ('Consume',   'Credit deducted; OpenAI call executed via the centralised AIService.',           'AIService'),
    ('Respond',   'AI result displayed to the user.',                                                'Feature output'),
    ('Depleted',  'If credits are exhausted, user is directed to upgrade their plan.',               '/pricing'),
])

sub_title(doc, '5.4  Authentication & Security')
flow_table(doc, [
    ('Login',       'User submits credentials — rate-limited to 5 attempts per minute.',            'throttle:5,1'),
    ('2FA',         'If enabled, user enters a TOTP code to complete authentication.',               '/two-factor-authentication'),
    ('Social Login','User authenticates via Google or LinkedIn OAuth 2.0.',                         'SocialAuthController'),
    ('Password Reset','Requests a reset link — rate-limited to 5 requests per hour.',               'Email token'),
    ('Session',     'Authenticated session established; all routes protected by middleware.',        'Auth guards'),
])

# Footer
para(doc, '')
divider(doc)
para(doc,
    'StudAI Hire  ·  User Journey Document  ·  Confidential  ·  May 28, 2026',
    size=8, colour=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, before=2, after=0)

output_path = r'C:\Users\user\Downloads\studai-career\StudAI_Hire_User_Journey_v3.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
